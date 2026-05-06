import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_table_from_data(doc, data):
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(data[0]):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row_data in data[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def create_lld_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Low Level Design (LLD) Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('CuraBot: Agentic AI Differential Diagnosis Tutor\nBased on Enterprise GenAI / Agentic AI Template')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1
    doc.add_heading('1. Document Control & Metadata', level=1)
    doc.add_paragraph('Purpose: Make the LLD auditable, versioned, and approved with clear AI safety classification.', style='Intense Quote')
    
    doc.add_heading('Document Header', level=2)
    add_table_from_data(doc, [
        ['Field', 'Details'],
        ['Project Name', 'CuraBot'],
        ['System Name', 'Agentic AI Differential Diagnosis Tutor'],
        ['Environment', 'Development / Educational'],
        ['Confidentiality', 'Internal / Restricted (Contains Clinical Logic)']
    ])
    
    doc.add_heading('Versioning & Change Log', level=2)
    add_table_from_data(doc, [
        ['Version', 'Date', 'Author', 'Description', 'Status'],
        ['v0.1', '2026-04-15', 'AI Architect', 'Initial Draft & SOP Framework', 'Draft'],
        ['v1.0', '2026-04-27', 'Project Team', 'Finalized Enterprise Template Formatting', 'Approved']
    ])
    
    doc.add_heading('Approval Matrix', level=2)
    add_table_from_data(doc, [
        ['Area', 'Reviewer', 'Approver', 'Evidence/Link'],
        ['Architecture', 'Lead Architect', 'ARB', 'ADR List / C4 Diagrams'],
        ['Security', 'Security Lead', 'CISO delegate', 'Threat model / SOP-010'],
        ['Compliance', 'Compliance SME', 'Compliance Board', 'Medical AI Policy Checklist'],
        ['Product', 'Product Owner', 'Product Head', 'BRD']
    ])
    
    doc.add_heading('AI Safety Classification & Autonomy', level=2)
    doc.add_paragraph('Safety Classification: PII/PHI Regulated (Handles simulated patient data)', style='List Bullet')
    doc.add_paragraph('Autonomy Level: L3 Supervised Agent (Recommend-only / Act-with-approval). Does not replace licensed medical practitioners.', style='List Bullet')
    
    doc.add_heading('Reference Links', level=2)
    doc.add_paragraph('BRD, HLD, ADRs, Threat Model, Automated Verification Suite Test Plan.', style='List Bullet')
    
    # Section 2
    doc.add_heading('2. System Overview', level=1)
    doc.add_paragraph('Purpose: Set clear scope boundaries so GenAI/agents do not accidentally expand behavior.', style='Intense Quote')
    
    doc.add_heading('Business Context & Problem Statement', level=2)
    doc.add_paragraph('Stakeholders: Medical Students (primary), Patients (primary), Curriculum Developers (consuming team).')
    doc.add_paragraph('Problem: Medical students lack structured tools to practice clinical reasoning outside the classroom. Patients struggle to understand symptom urgency.')
    
    doc.add_heading('Scope Boundaries', level=2)
    add_table_from_data(doc, [
        ['In-Scope User Journeys', 'Out-of-Scope (Non-Goals)'],
        ['Simulated differential diagnosis tutoring', 'Real-time clinical diagnosis for actual patients'],
        ['6-Agent reasoning pipeline execution', 'Prescribing medication or treatments'],
        ['Cognitive bias detection & feedback', 'Integration with live Hospital EMR systems']
    ])
    
    doc.add_heading('Success Metrics & KPIs', level=2)
    add_table_from_data(doc, [
        ['KPI', 'Target', 'Measurement Approach'],
        ['Diagnostic Alignment', '>90% match', 'Evaluation against 50 curated clinical cases'],
        ['Safety Adherence', '100% Red-Flag detection', 'Automated unit tests for Emergency SOPs'],
        ['Latency', '<3s response per agent step', 'Performance monitoring in dev console']
    ])
    
    doc.add_heading('Assumptions & Constraints', level=2)
    doc.add_paragraph('Assumptions: Reliable LLM API access (Gemini/Groq/OpenRouter).', style='List Bullet')
    doc.add_paragraph('Constraints: Limited to predefined 15 Clinical SOPs. Local compute constraints for React frontend.', style='List Bullet')
    
    # Section 3
    doc.add_heading('3. Architecture Decomposition', level=1)
    doc.add_paragraph('Purpose: Provide a layered breakdown (C4) plus deployment topology for enterprise runtime.', style='Intense Quote')
    
    doc.add_heading('C4 Context & Containers', level=2)
    doc.add_paragraph('Actors: Medical Students, Patients.')
    doc.add_paragraph('UI Container: React + TypeScript + Vite Single Page Application.')
    doc.add_paragraph('API Gateway & Orchestrator: FastAPI application running on Uvicorn.')
    doc.add_paragraph('Data Stores: Supabase (PostgreSQL BaaS) for user profiles/sessions, ChromaDB for vector embeddings.')
    
    doc.add_heading('C4 Components', level=2)
    doc.add_paragraph('Modules: LLMClient (routing), Orchestrator (workflow graph), PatientHistoryAnalyzer, MedicalCitationEngine, 15 SOP rule engines.')
    
    doc.add_heading('Deployment Topology', level=2)
    doc.add_paragraph('Frontend: Vercel/Netlify or equivalent static hosting.')
    doc.add_paragraph('Backend & Vector DB: Docker containerized on cloud VM / K8s cluster.')
    doc.add_paragraph('Environment Separation: Dev/Prod with separate Supabase projects and API keys via environment secrets.')
    
    # Section 4
    doc.add_heading('4. GenAI Capability Design', level=1)
    doc.add_paragraph('Purpose: Design RAG pipelines, prompt lifecycle, and hallucination controls.', style='Intense Quote')
    
    doc.add_heading('Retrieval Approach', level=2)
    doc.add_paragraph('Type: Vector-RAG via ChromaDB for medical report PDFs.')
    doc.add_paragraph('Scoping Rules: Enforced tenant/user separation (queries strictly scoped by user_id).')
    doc.add_paragraph('Chunking & Embedding: 600-character chunks to preserve clinical lines, embedded using GoogleGenerativeAiEmbeddingFunction.')
    
    doc.add_heading('Prompt Lifecycle', level=2)
    doc.add_paragraph('System Prompts: Static roles + behavioral constraints + SOPs to follow (e.g., OPQRST workflow).')
    doc.add_paragraph('Task Prompts: Dynamic templates injected with patient message, hypothesis state, and evidence ledger.')
    doc.add_paragraph('Output Formatting: Enforced JSON schema expected by downstream agents.')
    
    doc.add_heading('Guardrails & Confidence Scoring', level=2)
    doc.add_paragraph('Fallback Mechanism: Rule-based fallbacks for Agents 3, 4, and 6 to mitigate LLM failure/hallucination.')
    doc.add_paragraph('Grounding Policy: MedicalCitationEngine deterministically links hypotheses to ICD-10 and PubMed guidelines (zero hallucination risk).')
    doc.add_paragraph('Confidence Scoring: Bayesian likelihood ratio reasoning (Agent 4) with strict severity floors for critical diseases.')
    
    # Section 5
    doc.add_heading('5. Agentic Design', level=1)
    doc.add_paragraph('Purpose: Specify agent inventory, roles/skills, orchestration graph, and HITL checkpoints.', style='Intense Quote')
    
    doc.add_heading('Agent Inventory Table', level=2)
    add_table_from_data(doc, [
        ['Agent', 'Type', 'Inputs', 'Outputs', 'Tools Allowed', 'Cannot Do'],
        ['Symptom Normalizer', 'Extractor', 'Raw text + History', 'Normalized JSON', 'LLM / Keyword fallback', 'Generate diagnosis'],
        ['Hypothesis Gen.', 'Analyzer', 'Symptoms + KB', 'Ranked Hypotheses', 'Disease KB Search', 'Evaluate evidence'],
        ['Evidence Evaluator', 'Governance', 'Hypotheses + Evidence', 'Evidence Ledger', 'Rule-based matching', 'Update confidence'],
        ['Hypothesis Reviser', 'Math/Logic', 'Ledger', 'Revised Hypotheses', 'Bayesian multiplier', 'Ask questions'],
        ['Diagnostic Strategist', 'Decision', 'Revised Hypotheses', 'Next Q / Conclusion', 'LLM', 'Prescribe meds'],
        ['Self-Critique', 'Auditor', 'Hypotheses + State', 'Bias Flags JSON', 'Rule-based logic', 'Alter diagnosis']
    ])
    
    doc.add_heading('Orchestration Graph & Communication', level=2)
    doc.add_paragraph('Orchestration: LangGraph / DAG-based sequential execution.')
    doc.add_paragraph('Communication: Shared memory via the DiagnosticState Pydantic model (event bus pattern).')
    
    doc.add_heading('Human-in-the-Loop (HITL) Checkpoints', level=2)
    doc.add_paragraph('SOP-010 (Red Flag) & SOP-008 (Stroke): Act as immediate HITL/Emergency alerts, circumventing the LLM when critical conditions are flagged.')
    doc.add_paragraph('End of Session: Diagnostic report generation serves as an evidence pack for mentor review.')

    output_path = os.path.join(os.getcwd(), 'CuraBot_Enterprise_LLD.docx')
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == '__main__':
    create_lld_doc()
