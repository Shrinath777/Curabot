import sys
import subprocess
import os
import re

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Inches

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if in_table and not line.startswith('|'):
            # Render table
            if len(table_data) > 2: # Exclude the separator
                cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data)-1, cols=cols)
                table.style = 'Table Grid'
                # Header
                hdr_cells = table.rows[0].cells
                for i, cell_text in enumerate(table_data[0]):
                    if i < cols:
                        hdr_cells[i].text = cell_text.replace('**', '')
                # Rows
                row_idx = 1
                for r in table_data[2:]:
                    row_cells = table.rows[row_idx].cells
                    for i, cell_text in enumerate(r):
                        if i < cols:
                            row_cells[i].text = cell_text.replace('**', '').replace('<br/>', '\n').replace('<br>', '\n')
                    row_idx += 1
            in_table = False
            table_data = []

        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('![') and '](' in line:
            match = re.search(r'\!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"[Image placeholder: {img_path}]")
        elif line.startswith('|'):
            in_table = True
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            table_data.append(row)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
        elif line.startswith('> '):
            doc.add_paragraph(line[2:].replace('**', ''), style='Quote')
        elif line.startswith('```'):
            pass # ignore code block markers
        else:
            if not in_table and line != '---':
                doc.add_paragraph(line.replace('**', ''))

    if in_table and len(table_data) > 2:
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data)-1, cols=cols)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, cell_text in enumerate(table_data[0]):
            if i < cols:
                hdr_cells[i].text = cell_text.replace('**', '')
        row_idx = 1
        for r in table_data[2:]:
            row_cells = table.rows[row_idx].cells
            for i, cell_text in enumerate(r):
                if i < cols:
                    row_cells[i].text = cell_text.replace('**', '').replace('<br/>', '\n').replace('<br>', '\n')
            row_idx += 1

    doc.save(docx_path)

if __name__ == '__main__':
    md_path = r'c:\projects\tcs project\curabot\docs\LLD_CuraBot_Implementation_v1.md'
    docx_path = r'c:\projects\tcs project\curabot\docs\LLD_CuraBot_Implementation_v1.docx'
    md_to_docx(md_path, docx_path)
    print(f"Word document saved successfully at {docx_path}")
