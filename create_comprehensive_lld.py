import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_detailed_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

def create_comprehensive_lld():
    doc = Document()
    
    # --- TITLE PAGE ---
    title = doc.add_heading('Low Level Design (LLD) Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Enterprise GenAI / Agentic AI Application\nProject: CuraBot — Agentic AI Differential Diagnosis Tutor')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_page_break()
    
    # --- 1. Document Control & Metadata ---
    doc.add_heading('1. Document Control & Metadata', level=1)
    doc.add_paragraph('Purpose: Make the LLD auditable, versioned, and approved with clear AI safety classification.', style='Intense Quote')
    
    doc.add_heading('1.1 Document Header', level=2)
    add_detailed_table(doc, ['Field', 'Details'], [
        ['Project Name', 'CuraBot'],
        ['System Name', 'Agentic AI Differential Diagnosis Tutor'],
        ['Environment', 'Development / Educational (Targeted for Clinical Sandbox)'],
        ['Confidentiality', 'Internal / Restricted (Contains Proprietary Clinical Logic & Prompts)']
    ])
    
    doc.add_heading('1.2 Versioning & Change Log', level=2)
    add_detailed_table(doc, ['Version', 'Date', 'Author', 'Description', 'Status'], [
        ['v0.1', '2026-04-15', 'AI Architect', 'Initial Draft & SOP Framework Definition', 'Draft'],
        ['v0.5', '2026-04-20', 'Security Lead', 'Added Threat Model & Agent Guardrails', 'In Review'],
        ['v1.0', '2026-04-27', 'Project Team', 'Finalized Comprehensive GenAI Enterprise Format', 'Approved']
    ])
    
    doc.add_heading('1.3 Approval Matrix', level=2)
    add_detailed_table(doc, ['Area', 'Reviewer', 'Approver', 'Evidence/Link'], [
        ['Architecture', 'Lead Architect', 'ARB / TRB', 'ADR List / Orchestration Graph'],
        ['Security', 'Security Lead', 'CISO Delegate', 'Threat Model / 10 Clinical SOPs Check'],
        ['Compliance', 'Compliance SME', 'Compliance Board', 'Medical AI & PII Policy Checklist'],
        ['Product', 'Product Owner', 'Product Head', 'BRD & KPI Alignment Matrix'],
        ['Operations', 'DevOps Lead', 'Infrastructure Head', 'Deployment Topology / Resource Specs']
    ])
    
    doc.add_heading('1.4 AI Safety Classification & Data Handling', level=2)
    doc.add_paragraph('Classification:', style='List Bullet')
    doc.add_paragraph('PII/PHI Regulated. The system actively processes simulated patient symptoms, clinical hypotheses, and uploaded medical PDF reports.', style='List Bullet 2')
    doc.add_paragraph('Data Handling Summary:', style='List Bullet')
    doc.add_paragraph('All data is anonymized before LLM transmission. No real-world PII should be ingested into the educational LLM prompts. RAG context is strictly isolated per user session (tenant scoping).', style='List Bullet 2')
    
    doc.add_heading('1.5 AI Autonomy Level', level=2)
    doc.add_paragraph('Autonomy Level: L3 — Supervised Agent ("Act-with-approval" / "Recommend-only").')
    doc.add_paragraph('The system acts as a conversational diagnostician that recommends conditions and requests human feedback/approval. It NEVER prescribes medication, executes external treatments, or updates live EMR records autonomously. All actions are overseen by the human user.')
    
    doc.add_heading('1.6 Reference Artifacts', level=2)
    doc.add_paragraph('BRD: Business Requirements Document detailing user journeys.', style='List Bullet')
    doc.add_paragraph('HLD: High-Level Design / Cloud Architecture Diagram.', style='List Bullet')
    doc.add_paragraph('ADRs: Architectural Decision Records (e.g., SQLite to Supabase migration, LLM routing logic).', style='List Bullet')
    doc.add_paragraph('Test Plan: Automated Verification Suite for all 15 Clinical SOPs.', style='List Bullet')
    
    # --- 2. System Overview ---
    doc.add_heading('2. System Overview', level=1)
    doc.add_paragraph('Purpose: Set clear scope boundaries so GenAI/agents do not "accidentally" expand behavior.', style='Intense Quote')
    
    doc.add_heading('2.1 Business Context & Stakeholders', level=2)
    doc.add_paragraph('Primary Personas:', style='List Bullet')
    doc.add_paragraph('Medical Students: Practicing differential diagnosis by observing AI reasoning, evaluating evidence, and understanding confidence trajectories over multiple questioning rounds.', style='List Bullet 2')
    doc.add_paragraph('General Users (Patients): Gaining preliminary understanding of their symptoms and urgency (Triage) before consulting a licensed doctor.', style='List Bullet 2')
    doc.add_paragraph('Consuming Teams:', style='List Bullet')
    doc.add_paragraph('Curriculum Developers: Integrating the bot into clinical simulations and using performance logs for grading.', style='List Bullet 2')
    
    doc.add_heading('2.2 Problem Statement', level=2)
    doc.add_paragraph('Medical students and early-career healthcare learners often lack access to structured, dynamic tools for practicing clinical reasoning outside the classroom. Textbooks are static, and clinical rotations are limited.')
    doc.add_paragraph('Measurable Pain: High manual effort for research, latency in feedback loops, and high risk of clinical cognitive biases (e.g., anchoring or premature closure). Patients face delayed care due to a lack of immediate symptom triage.')
    
    doc.add_heading('2.3 Scope Boundaries', level=2)
    add_detailed_table(doc, ['In-Scope User Journeys', 'Out-of-Scope (Explicit Non-Goals)'], [
        ['Symptom Normalization from plain text', 'Real-time diagnosis of actual/live patients'],
        ['Generation of 3-5 Ranked Differential Diagnoses', 'Prescribing medication, procedures, or treatments'],
        ['Multi-turn conversational follow-up questions', 'Integration with live Hospital EMR / EHR systems'],
        ['PDF Medical Report Upload & RAG processing', 'Replacing licensed medical practitioners'],
        ['Emergency Triage & Red-Flag Interception', 'Autonomous booking of specialist appointments'],
        ['Detection of cognitive biases in reasoning', 'Long-term patient health management/tracking']
    ])
    
    doc.add_heading('2.4 Success Metrics & KPIs', level=2)
    add_detailed_table(doc, ['KPI', 'Target', 'Baseline', 'Measurement Approach'], [
        ['Diagnostic Alignment', '>90% match w/ Gold Standard', '65% (Static tools)', 'Automated evaluation against 50 curated clinical cases'],
        ['Safety Adherence', '100% Red-Flag detection', '80%', 'Unit tests targeting SOP-013 & SOP-008'],
        ['User Engagement', '>15 mins average session', '5 mins', 'WebSocket / Supabase session activity logging'],
        ['System Latency', '<3s response per agent step', '10s', 'Telemetry and Dev Console monitoring'],
        ['RAG Accuracy', '>95% precision on lab extraction', '70%', 'Ground-truth testing on synthetic PDF reports']
    ])
    
    doc.add_heading('2.5 Assumptions & Constraints', level=2)
    doc.add_paragraph('Assumptions: Reliable external API access for LLMs (Gemini, Groq, OpenRouter). User possesses basic understanding of their symptoms.', style='List Bullet')
    doc.add_paragraph('Constraints: System strictly constrained to the 15 implemented Clinical SOPs. Local compute limitations affect Vector DB processing (ChromaDB runs locally).', style='List Bullet')
    
    # --- 3. Architecture Decomposition ---
    doc.add_heading('3. Architecture Decomposition', level=1)
    doc.add_paragraph('Purpose: Provide a layered breakdown (C4) plus deployment topology for enterprise runtime.', style='Intense Quote')
    
    doc.add_heading('3.1 C4-Context', level=2)
    doc.add_paragraph('Actors: End Users (Students, Patients), Mentors/Reviewers.')
    doc.add_paragraph('Downstream Systems: Supabase (PostgreSQL), ChromaDB (Vector DB), LLM APIs (Gemini/Groq/OpenRouter).')
    doc.add_paragraph('Trust Boundaries: All client communication terminates at the FastAPI API Gateway. The agent runtime and orchestration occur securely backend-side. No direct database access from the frontend.')
    
    doc.add_heading('3.2 C4-Container', level=2)
    add_detailed_table(doc, ['Container', 'Tech Stack', 'Responsibility'], [
        ['User Interface (UI)', 'React, TypeScript, Vite, Zustand', 'Renders chat, hypothesis cards, trajectory charts.'],
        ['API Gateway', 'FastAPI, Uvicorn, Python', 'Handles REST endpoints (/api/chat, /api/auth).'],
        ['Orchestrator', 'LangGraph Core', 'Coordinates the 6-agent pipeline workflow.'],
        ['Agent Runtime', 'Python Agents (Classes)', 'Executes individual tasks (Normalization, Eval).'],
        ['Data Stores', 'Supabase (Cloud) / SQLite (Local)', 'Stores users, sessions, profiles, chat history.'],
        ['Vector Database', 'ChromaDB', 'Stores embeddings of medical PDFs for RAG.']
    ])
    
    doc.add_heading('3.3 C4-Component', level=2)
    doc.add_paragraph('Key Modules:', style='List Bullet')
    doc.add_paragraph('LLMClient (services/llm_client.py): Manages provider fallback chain (Gemini -> Groq -> OpenRouter) and applies a 0.5s pacing lock to mitigate rate limits.', style='List Bullet 2')
    doc.add_paragraph('RAG Retriever (ChromaDB): Extracts PDF via pdfplumber, chunks at 600 characters, indexes, and retrieves top-3 relevant chunks per user query.', style='List Bullet 2')
    doc.add_paragraph('Medical Evidence Citation Engine: Pure Python module generating ICD-10 codes, Clinical Authority links (AHA, CDC), and PubMed verification links for zero-hallucination grounding.', style='List Bullet 2')
    doc.add_paragraph('Clinical SOP Engine: 15 rule-based engines running pre/post LLM to enforce safety (Triage, Stroke, Lab interpretation).', style='List Bullet 2')
    
    doc.add_heading('3.4 Deployment Topology & Environment', level=2)
    doc.add_paragraph('Frontend Deployment: Hosted on Vercel / Netlify edge networks.')
    doc.add_paragraph('Backend & Vector DB Deployment: Containerized using Docker. Deployed to Kubernetes (K8s) or Cloud VMs with isolated Virtual Private Cloud (VPC) subnets.')
    doc.add_paragraph('Environment Separation: Strict isolation between Dev, Staging (UAT), and Prod. Supabase environments are totally distinct. Secrets and API keys are injected via secure Environment Variables (.env config strategy).')
    
    # --- 4. GenAI Capability Design ---
    doc.add_heading('4. GenAI Capability Design', level=1)
    doc.add_paragraph('Purpose: Design RAG pipelines, prompt lifecycle, and hallucination controls.', style='Intense Quote')
    
    doc.add_heading('4.1 Retrieval Approach (Vector-RAG)', level=2)
    doc.add_paragraph('Approach: Vector-RAG applied to the patient_records collection for uploaded medical lab reports and PDFs.')
    doc.add_paragraph('Scoping Rules: Queries are strictly scoped by user_id to ensure tenants (patients) can never retrieve medical data belonging to another user. Access control is enforced at the API route layer.')
    doc.add_paragraph('Chunking Strategy: Text is extracted using pdfplumber, split into 600-character chunks (preserving line boundaries for clinical parameters like "Hb: 12.5 g/dL").')
    doc.add_paragraph('Embedding Model: GoogleGenerativeAiEmbeddingFunction when available, with a zero-vector local fallback if the API is offline.')
    
    doc.add_heading('4.2 Prompt Lifecycle', level=2)
    doc.add_paragraph('System Prompts (Agent Identity): Static blocks defining roles, clinical rules, and SOPs. Example: Diagnostic Strategist is instructed to follow the OPQRST clinical workflow.', style='List Bullet')
    doc.add_paragraph('Task Prompts (Dynamic Context): Template-based prompts filled with the current DiagnosticState (patient message, hypotheses, evidence ledger, RAG context, and iteration count).', style='List Bullet')
    doc.add_paragraph('Evaluation Prompts: Used by the Patient History Analyzer to compare current symptoms against recurring conditions stored in Supabase.', style='List Bullet')
    
    doc.add_heading('4.3 Guardrails & Output Validation', level=2)
    doc.add_paragraph('Schema Validation: All LLM outputs are mandated to follow exact JSON schema structures. Fallback parsers handle markdown-wrapped or malformed outputs.', style='List Bullet')
    doc.add_paragraph('Prompt Injection Mitigation: Raw user input is constrained. Emergency phrases are intercepted by the pure-Python SOP-013 (Red Flag Scanner) BEFORE reaching the LLM.', style='List Bullet')
    doc.add_paragraph('Citation Requirement: The Medical Evidence Citation Engine forces every concluded diagnosis to be backed by deterministically mapped ICD-10 codes and PubMed links (Grounding Policy).', style='List Bullet')
    
    doc.add_heading('4.4 Confidence Scoring Approach', level=2)
    doc.add_paragraph('Confidence scores are dynamically updated by Agent 4 (Hypothesis Reviser) using a Bayesian likelihood ratio approach:')
    doc.add_paragraph('• Strong supporting evidence multiplies confidence by 1.6x.\n• Contradicting evidence applies reduction multipliers (0.5x, 0.7x).\n• Severity Floors: Critical-severity diseases cannot drop below 8% unless they have 3+ contradicting evidence items, ensuring they are never prematurely dismissed.')
    
    # --- 5. Agentic Design ---
    doc.add_heading('5. Agentic Design', level=1)
    doc.add_paragraph('Purpose: Specify agent inventory, roles/skills, orchestration graph, and HITL checkpoints.', style='Intense Quote')
    
    doc.add_heading('5.1 Agent Inventory Table', level=2)
    add_detailed_table(doc, ['Agent', 'Type', 'Inputs', 'Outputs', 'Tools Allowed', 'Cannot Do'], [
        ['Symptom Normalizer', 'Extractor', 'Raw text + RAG History', 'Normalized JSON', 'LLM, Keyword Matching', 'Generate diagnosis, Alter context'],
        ['Hypothesis Generator', 'Analyzer', 'Symptoms + Disease KB', 'Ranked Hypotheses', 'KB Search, LLM', 'Evaluate evidence, Ask user'],
        ['Evidence Evaluator', 'Governance', 'Hypotheses + Evidence', 'Evidence Ledger', 'Rule-based evaluation', 'Update confidence scores'],
        ['Hypothesis Reviser', 'Math/Logic', 'Ledger + Hypotheses', 'Revised Hypotheses', 'Bayesian Multipliers', 'Ask questions, Access LLM'],
        ['Diagnostic Strategist', 'Decision', 'Revised Hypotheses', 'Next Q / Conclusion', 'LLM, OPQRST Workflow', 'Prescribe medications'],
        ['Self-Critique', 'Auditor', 'State + Iteration count', 'Bias Flags (JSON)', 'Rule-based checking', 'Alter the final diagnosis']
    ])
    
    doc.add_heading('5.2 Orchestration Graph & Communication', level=2)
    doc.add_paragraph('Graph Execution: Managed as a sequential state machine (Directed Acyclic Graph) via LangGraph principles. Pipeline runs once per user message (typically completing in 3-10 iterations).')
    doc.add_paragraph('Communication Pattern: Agents communicate via Shared Memory. The entire session context is persisted in the DiagnosticState Pydantic model (acting as the state payload passed from node to node).')
    doc.add_paragraph('Cost & Latency Optimization: Only Agents 1, 2, and 5 invoke the LLM. Agents 3, 4, and 6 use high-performance Python rule-based logic to save API costs and reduce latency.')
    
    doc.add_heading('5.3 Human-in-the-Loop (HITL) Checkpoints', level=2)
    doc.add_paragraph('SOP-010 (Red Flag Scanner) & SOP-008 (FAST Stroke): These Python protocols run immediately after normalization. If a stroke indicator or critical red flag is detected, the orchestrator immediately halts the LLM pipeline and triggers an Emergency Override Alert to the user to seek immediate care.', style='List Bullet')
    doc.add_paragraph('Diagnostic Report Conclusion: The final output is flagged clearly as "For Medical Education Only". Any diagnostic application must be reviewed by a human mentor or licensed physician. The system generates an evidence pack (JSON/HTML) for this explicit HITL review.', style='List Bullet')
    
    doc.add_heading('5.4 Release & Versioning Strategy', level=2)
    doc.add_paragraph('Workflow Versioning: The orchestration graph is versioned in Git. Changes to agent sequences require regression testing against the 15 Automated SOP Tests.')
    doc.add_paragraph('Prompt Pack Versioning: System prompts are stored as constants and versioned alongside the codebase. Changes to clinical guidelines (e.g., AHA updates) require a major version bump in the prompt registry.')
    
    # Save the document
    output_path = os.path.join(os.getcwd(), 'CuraBot_Comprehensive_LLD.docx')
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == '__main__':
    create_comprehensive_lld()
