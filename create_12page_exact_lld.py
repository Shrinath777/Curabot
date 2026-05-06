import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_detailed_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

def create_12page_exact_lld():
    doc = Document()
    
    # ---------------------------------------------------------
    # TITLE PAGE (PAGE 1)
    # ---------------------------------------------------------
    for _ in range(5): doc.add_paragraph()
    title = doc.add_heading('Low Level Design (LLD) Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Enterprise GenAI / Agentic AI Application\nProject: CuraBot — Agentic AI Differential Diagnosis Tutor\nVersion 1.0 (Comprehensive Build)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(15): doc.add_paragraph()
    doc.add_paragraph('Disclaimer: FOR MEDICAL EDUCATION ONLY — Not for clinical use', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 1. DOCUMENT CONTROL & METADATA (PAGE 2)
    # ---------------------------------------------------------
    doc.add_heading('1. Document Control & Metadata', level=1)
    doc.add_paragraph('Purpose: Make the LLD auditable, versioned, and approved with clear AI safety classification.', style='Intense Quote')
    
    doc.add_heading('1.1 Document Header', level=2)
    add_detailed_table(doc, ['Project', 'System Name', 'Environment', 'Confidentiality'], [
        ['CuraBot', 'Agentic AI Differential Diagnosis Tutor', 'Dev / Int / UAT / Prod (Clinical Sandbox)', 'Strictly Internal / Restricted']
    ])
    doc.add_paragraph()
    
    doc.add_heading('1.2 Doc Control Table (Versioning & Change Log)', level=2)
    add_detailed_table(doc, ['Version', 'Author', 'Reviewer', 'Approver', 'Date', 'Change Log'], [
        ['v0.1 draft', 'AI Architect', 'Peer Architect', 'Lead Architect', '2026-04-10', 'Initial Draft'],
        ['v0.5', 'Security Lead', 'InfoSec Reviewer', 'CISO', '2026-04-15', 'Added Threat Model'],
        ['v1.0 approved', 'Project Team', 'Security Lead', 'ARB/TRB', '2026-04-27', 'Final Comprehensive Approval']
    ])
    doc.add_paragraph()
    
    doc.add_heading('1.3 Approval Matrix', level=2)
    add_detailed_table(doc, ['Area', 'Reviewer', 'Approver', 'Evidence/Link'], [
        ['Architecture', 'Lead Architect', 'ARB/TRB', 'ADR list & C4 Diagrams'],
        ['Security', 'Security Lead', 'CISO delegate', 'Threat model & 15 SOPs'],
        ['Compliance', 'Compliance SME', 'AIRB/Compliance board', 'Medical AI Policy checklist'],
        ['Product', 'Product Owner', 'VP of Product', 'BRD Alignment']
    ])
    doc.add_paragraph()
    
    doc.add_heading('1.4 AI Safety Classification & Data Handling Summary', level=2)
    doc.add_paragraph('Safety Classification:', style='List Bullet')
    doc.add_paragraph('PII/PHI Regulated. The system processes simulated patient symptoms and PDF medical reports.', style='List Bullet 2')
    doc.add_paragraph('Data Sensitivity Summary:', style='List Bullet')
    doc.add_paragraph('High. All data is anonymized prior to any LLM transmission. No real-world PII is ingested into the educational LLM prompts. RAG context is strictly isolated per user session (tenant scoping) utilizing ChromaDB namespaces. The system is designed purely for medical education, not live patient treatment.', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('1.5 AI Autonomy Level', level=2)
    doc.add_paragraph('Autonomy Level: L3 — Supervised Agent ("Act-with-approval" / "Recommend-only").')
    doc.add_paragraph('The system acts as a conversational diagnostician that recommends conditions and requests human feedback/approval. It NEVER prescribes medication, executes external treatments, or updates live EMR records autonomously. All actions are overseen by the human user.')
    doc.add_paragraph()
    
    doc.add_heading('1.6 Reference Artifacts', level=2)
    doc.add_paragraph('BRD: Business Requirements Document detailing user journeys.', style='List Bullet')
    doc.add_paragraph('HLD: High-Level Design / Cloud Architecture Diagram.', style='List Bullet')
    doc.add_paragraph('ADRs: Architectural Decision Records (e.g., Supabase migration, LLM routing logic).', style='List Bullet')
    doc.add_paragraph('Test Plan: Automated Verification Suite for all 15 Clinical SOPs.', style='List Bullet')
    doc.add_paragraph('Threat Model: Agentic pipeline injection mitigation strategies.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 2. SYSTEM OVERVIEW (PAGE 3)
    # ---------------------------------------------------------
    doc.add_heading('2. System Overview', level=1)
    doc.add_paragraph('Purpose: Set clear scope boundaries so GenAI/agents do not "accidentally" expand behavior.', style='Intense Quote')
    
    doc.add_heading('2.1 Business Context, Stakeholders & Personas', level=2)
    doc.add_paragraph('Medical students and early-career healthcare learners often lack access to structured tools that can help them practice clinical reasoning skills outside the classroom.')
    doc.add_paragraph('Primary Personas:', style='List Bullet')
    doc.add_paragraph('Medical Students: The core audience practicing differential diagnosis by entering symptoms and observing AI reasoning over multiple rounds of questioning.', style='List Bullet 2')
    doc.add_paragraph('General Users (Patients): Individuals wanting a preliminary understanding of their symptoms and triage urgency.', style='List Bullet 2')
    doc.add_paragraph('Consuming Teams:', style='List Bullet')
    doc.add_paragraph('Curriculum Developers: Integrating the bot into clinical simulations and using logs for grading.', style='List Bullet 2')
    doc.add_paragraph('Clinical Mentors: Providing human-in-the-loop oversight on generated diagnostic reports.', style='List Bullet 2')
    doc.add_paragraph('Primary Use Cases: 1) Multi-turn diagnostic reasoning, 2) Triage Red-Flag scanning, 3) Medical PDF RAG interpretation, 4) Bayesian confidence tracking.')
    doc.add_paragraph()
    
    doc.add_heading('2.2 Problem Statement & Measurable Pain', level=2)
    doc.add_paragraph('The lack of dynamic, accessible clinical reasoning tools leads to significant skill gaps in diagnostic formulation.')
    doc.add_paragraph('Measurable Pain Points:', style='List Bullet')
    doc.add_paragraph('High manual effort required by faculty to create and grade synthetic clinical cases.', style='List Bullet 2')
    doc.add_paragraph('Latency in feedback loops; students wait weeks to receive feedback on diagnostic assessments.', style='List Bullet 2')
    doc.add_paragraph('High risk of clinical cognitive biases (e.g., anchoring or premature closure) going uncorrected during formative training years.', style='List Bullet 2')
    doc.add_paragraph('Patients face delayed care due to a lack of immediate, accessible symptom triage.', style='List Bullet 2')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 2. SYSTEM OVERVIEW CONT. (PAGE 4)
    # ---------------------------------------------------------
    doc.add_heading('2.3 Scope Boundaries (In-Scope / Out-of-Scope)', level=2)
    add_detailed_table(doc, ['In-Scope User Journeys', 'Out-of-Scope (Explicit Non-Goals)'], [
        ['Symptom Normalization from plain text', 'Real-time diagnosis of actual/live patients'],
        ['Generation of 3-5 Ranked Differential Diagnoses', 'Prescribing medication, procedures, or treatments'],
        ['Multi-turn conversational follow-up questions', 'Integration with live Hospital EMR / EHR systems'],
        ['PDF Medical Report Upload & RAG processing', 'Replacing licensed medical practitioners'],
        ['Emergency Triage & Red-Flag Interception', 'Autonomous booking of specialist appointments'],
        ['Detection of cognitive biases in reasoning', 'Long-term patient health management/tracking'],
        ['Bayesian confidence level updates over time', 'Integration with wearable health devices (IoT)'],
        ['Profile management (History, Allergies, Meds)', 'Automated billing or medical coding']
    ])
    doc.add_paragraph()
    
    doc.add_heading('2.4 Success Metrics & KPIs', level=2)
    doc.add_paragraph('To ensure the agentic system delivers value without compromising clinical education safety, the following strict Service Level Agreements (SLAs) and KPIs are monitored.')
    add_detailed_table(doc, ['KPI Category', 'Measurement Approach (Baseline -> Target)'], [
        ['Quality (Diagnostic Alignment)', 'Automated case evaluation (65% -> >90% match)'],
        ['Time (Latency)', 'Dev Console telemetry (10s -> <3s per step)'],
        ['Compliance (Safety Adherence)', 'SOP Unit tests (80% -> 100% Red-Flag detection)'],
        ['Adoption (Engagement)', 'Session logging (5 mins -> >15 mins)'],
        ['RAG Accuracy', 'Ground-truth extraction testing (70% -> >95% precision)']
    ])
    doc.add_paragraph()
    
    doc.add_heading('2.5 Assumptions, Constraints & Autonomy Target', level=2)
    doc.add_paragraph('Assumptions: Data availability of disease knowledge base, LLM model access via API. Users can communicate symptoms effectively.', style='List Bullet')
    doc.add_paragraph('Constraints: Tool access is strictly limited to the 15 defined SOPs. No external live web scraping allowed. Local compute limitations affect Vector DB processing.', style='List Bullet')
    doc.add_paragraph('Autonomy Level Target: Started as L1 assistant → final architecture target achieved is L3 Supervised Agent (Recommended for regulated environments).', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 3. ARCHITECTURE DECOMPOSITION (PAGE 5)
    # ---------------------------------------------------------
    doc.add_heading('3. Architecture Decomposition', level=1)
    doc.add_paragraph('Purpose: Provide a layered breakdown (C4) plus deployment topology for enterprise runtime.', style='Intense Quote')
    
    doc.add_heading('3.1 C4-Context', level=2)
    doc.add_paragraph('Actors: End Users (Medical Students, Patients), Clinical Mentors / Reviewers.')
    doc.add_paragraph('Upstream Systems: None (System is the primary entry point).')
    doc.add_paragraph('Downstream Systems: Supabase (PostgreSQL BaaS), ChromaDB (Vector DB), LLM API Providers (Gemini, Groq, OpenRouter).')
    doc.add_paragraph('Trust Boundaries: All client communication terminates at the FastAPI API Gateway over HTTPS. Agent runtime is secured behind the backend firewall.')
    doc.add_paragraph('\n[PLACEHOLDER: Insert PlantUML / draw.io C4-Context Diagram Here]\n')
    
    doc.add_heading('3.2 C4-Container', level=2)
    doc.add_paragraph('The system is distributed across several decoupled containers:')
    add_detailed_table(doc, ['Container', 'Tech Stack', 'Responsibility'], [
        ['UI', 'React, TypeScript, Vite', 'Renders chat, hypothesis cards, trajectory charts.'],
        ['API Gateway', 'FastAPI, Uvicorn', 'Handles REST endpoints (/api/chat, /api/auth).'],
        ['Orchestrator', 'LangGraph Core', 'Coordinates the 6-agent pipeline workflow.'],
        ['Agent Runtime', 'Python Agents', 'Executes individual tasks (Normalization, Eval).'],
        ['Data Stores', 'Supabase / SQLite', 'Stores users, sessions, profiles, chat history.'],
        ['Retrieval', 'ChromaDB', 'Stores embeddings of medical PDFs for RAG.']
    ])
    doc.add_paragraph('\n[PLACEHOLDER: Insert PlantUML / draw.io C4-Container Diagram Here]\n')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 3. ARCHITECTURE DECOMPOSITION CONT. (PAGE 6)
    # ---------------------------------------------------------
    doc.add_heading('3.3 C4-Component: Component Inventory with Owners', level=2)
    doc.add_paragraph('Within the backend container, the architecture is highly modularized:')
    add_detailed_table(doc, ['Component', 'Description', 'Owner/Team'], [
        ['Retriever Module', 'ChromaDB interface for vector RAG ingestion/retrieval.', 'Data Eng Team'],
        ['Ranker / Router', 'LLMClient API rotation logic and rate limit locks.', 'Backend Core Team'],
        ['Prompt Builder', 'Dynamic DiagnosticState injector for Task Prompts.', 'AI/ML Team'],
        ['Validators (SOPs)', '15 pure-Python guardrail rule engines (e.g., Triage, Red Flag).', 'Clinical QA Team'],
        ['Citation Engine', 'PubMed & ICD-10 grounding linker.', 'Clinical QA Team']
    ])
    doc.add_paragraph()
    
    doc.add_heading('3.4 C4-Code', level=2)
    doc.add_paragraph('Key Packages/Classes:', style='List Bullet')
    doc.add_paragraph('`services/orchestrator.py`: Implements the LangGraph execution.', style='List Bullet 2')
    doc.add_paragraph('`services/llm_client.py`: Implements provider fallback (Gemini -> Groq -> OpenRouter).', style='List Bullet 2')
    doc.add_paragraph('Interface Contracts: Standardized REST JSON schemas for all frontend-backend communication.', style='List Bullet')
    doc.add_paragraph('Configuration Boundaries: Controlled exclusively via strict `.env` environmental variables.', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('3.5 Deployment Topology & Environments', level=2)
    doc.add_paragraph('Topology: Deployed via Docker into a K8s cluster. Namespaces strictly separate `curabot-ui` and `curabot-api`. Network zones/VNETs provide deep isolation. Private endpoints are utilized for all Supabase cloud database communication.')
    doc.add_paragraph('Environment Separation: Dev, Int, UAT, Prod.')
    doc.add_paragraph('Config Strategy: Secrets injected securely via Vault; feature flags (LaunchDarkly) used to toggle LLM routing.')
    doc.add_paragraph('\n[PLACEHOLDER: Insert Deployment diagram + network zones Here]\n')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 4. GENAI CAPABILITY DESIGN (PAGE 7)
    # ---------------------------------------------------------
    doc.add_heading('4. GenAI Capability Design', level=1)
    doc.add_paragraph('Purpose: Design RAG pipelines, prompt lifecycle, and hallucination controls.', style='Intense Quote')
    
    doc.add_heading('4.1 RAG Pipeline Flow & Retrieval Approach', level=2)
    doc.add_paragraph('Approach: Vector-RAG combined with specific metadata filtering.')
    doc.add_paragraph('RAG Pipeline Flow (Step-by-step):', style='List Bullet')
    doc.add_paragraph('1. INGEST: `pdfplumber` extracts readable text from user-uploaded Medical PDFs.', style='List Bullet 2')
    doc.add_paragraph('2. INDEX: Text is split into 600-character chunks and embedded using `GoogleGenerativeAiEmbeddingFunction`.', style='List Bullet 2')
    doc.add_paragraph('3. RETRIEVE: A cosine similarity search queries the database using the patient’s chat input.', style='List Bullet 2')
    doc.add_paragraph('4. RERANK: The Top-K (3) most relevant chunks are selected based on vector distance.', style='List Bullet 2')
    doc.add_paragraph('5. GENERATE: The retrieved chunks are injected directly into the Prompt Builder context window.', style='List Bullet 2')
    doc.add_paragraph('6. VALIDATE: The final LLM response is checked against strict JSON schema validators.', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('4.2 Document/Query Scoping Rules', level=2)
    doc.add_paragraph('Tenant Scoping: Queries are strictly scoped by `user_id` at the database level to ensure tenants (patients) can never retrieve medical data belonging to another user. Access control is enforced via Supabase Row Level Security (RLS) and API route verification.')
    doc.add_paragraph('Chunking Strategy: Text is split into exactly 600-character chunks to preserve clinical line boundaries (e.g., keeping lab values attached to their test names without truncation).')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 4. GENAI CAPABILITY DESIGN CONT. (PAGE 8)
    # ---------------------------------------------------------
    doc.add_heading('4.3 Prompt Registry & Lifecycle', level=2)
    doc.add_paragraph('CuraBot uses structured, version-controlled prompt engineering to ensure consistent, parseable outputs.')
    add_detailed_table(doc, ['Prompt Type', 'Version', 'Approval Status', 'Purpose / Usage'], [
        ['System Prompts', 'v1.2', 'Approved', 'Static instructions defining agent roles & clinical rules (e.g., OPQRST workflow).'],
        ['Task Prompts', 'v1.0', 'Approved', 'Dynamic templates injected with DiagnosticState (symptoms, history, evidence).'],
        ['Retrieval Templates', 'v1.1', 'Approved', 'Formats Top-K chunks into structured "UPLOADED LABS" context blocks.'],
        ['Evaluation Prompts', 'v1.0', 'Approved', 'Used by Self-Critique Agent for Bias checking (e.g., anchoring bias).']
    ])
    doc.add_paragraph()
    
    doc.add_heading('4.4 Guardrails & Grounding Policy', level=2)
    doc.add_paragraph('Prompt-injection mitigation: Raw user input is constrained. Emergency phrases are intercepted by the pure-Python SOP-010 (Red Flag Scanner) BEFORE reaching the LLM.', style='List Bullet')
    doc.add_paragraph('Output schema validation: All LLM outputs must follow exact JSON schema structures. Downstream fallback parsers handle markdown-wrapped or malformed outputs.', style='List Bullet')
    doc.add_paragraph('Grounding Policy ("Answer only from retrieved evidence"): The Medical Evidence Citation Engine forces every concluded diagnosis to be backed by deterministically mapped ICD-10 codes and PubMed verification links. This ensures zero-hallucination grounding for all diagnoses.', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('4.5 Confidence Scoring Approach', level=2)
    doc.add_paragraph('Confidence scores are dynamically updated by Agent 4 (Hypothesis Reviser) using a Bayesian likelihood ratio approach:')
    doc.add_paragraph('Generation Confidence Adjustments:', style='List Bullet')
    doc.add_paragraph('Strong supporting evidence multiplies hypothesis confidence by 1.6x.', style='List Bullet 2')
    doc.add_paragraph('Moderate supporting evidence multiplies by 1.3x.', style='List Bullet 2')
    doc.add_paragraph('Contradicting evidence applies reduction multipliers (Strong: 0.5x, Moderate: 0.7x).', style='List Bullet 2')
    doc.add_paragraph('Thresholds (Severity Floors): Critical-severity diseases cannot drop below an 8% baseline unless they have 3+ contradicting evidence items, ensuring deadly conditions are never prematurely dismissed.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 5. AGENTIC DESIGN (PAGE 9)
    # ---------------------------------------------------------
    doc.add_heading('5. Agentic Design', level=1)
    doc.add_paragraph('Purpose: Specify agent inventory, roles/skills, orchestration graph, and HITL checkpoints.', style='Intense Quote')
    
    doc.add_heading('5.1 Agent Inventory Table & RACI Matrix', level=2)
    doc.add_paragraph('CuraBot operates a strict 6-agent sequential pipeline.')
    add_detailed_table(doc, ['Agent', 'Type', 'Inputs', 'Outputs', 'Tools Allowed', 'Cannot Do'], [
        ['1. Symptom Normalizer', 'Extractor', 'User Msg + RAG', 'Normalized Symptoms JSON', 'LLM, Keyword Fallback', 'Generate diagnoses, Alter context'],
        ['2. Hypothesis Gen.', 'Analyzer', 'Symptoms', 'Ranked Hypotheses JSON', 'Disease KB Search, LLM', 'Evaluate evidence, Ask final questions'],
        ['3. Evidence Evaluator', 'Governance', 'Hypotheses + Evidence', 'Evidence Ledger JSON', 'Rule-based evaluation', 'Update confidence scores'],
        ['4. Hypothesis Reviser', 'Math/Logic', 'Ledger + Hypotheses', 'Revised Hypotheses JSON', 'Bayesian Multipliers Tool', 'Ask questions, Access LLM'],
        ['5. Diagnostic Strategist', 'Decision', 'Revised Hypotheses', 'Next Q / Conclusion JSON', 'LLM, OPQRST Tool', 'Prescribe medications or procedures'],
        ['6. Self-Critique', 'Auditor', 'State + Iteration count', 'Bias Flags (JSON)', 'Rule-based auditing logic', 'Alter the final diagnosis']
    ])
    doc.add_paragraph()
    
    doc.add_heading('5.2 Clinical SOP Integrations', level=2)
    doc.add_paragraph('The agents are restricted by exactly 15 pure Python rule-based SOPs. Examples include:')
    doc.add_paragraph('SOP-001 Triage Severity Classification: Analyzes vitals to assign Red/Yellow/Green urgency.', style='List Bullet')
    doc.add_paragraph('SOP-002 Chest Pain Protocol: Detects crushing pain radiating to the jaw.', style='List Bullet')
    doc.add_paragraph('SOP-003 Stroke Detection (FAST): Identifies facial drooping or speech difficulty.', style='List Bullet')
    doc.add_paragraph('SOP-008 Lab Value Interpreter: Parses extracted PDF text for critical lab abnormalities.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 5. AGENTIC DESIGN CONT. (PAGE 10)
    # ---------------------------------------------------------
    doc.add_heading('5.3 Orchestration Graph & State Machine', level=2)
    doc.add_paragraph('Design: Managed as a sequential state machine (Directed Acyclic Graph) via LangGraph principles. The pipeline runs exactly once per user message and typically completes the diagnostic cycle in 3-10 iterations.')
    doc.add_paragraph('States & Transitions: State transitions rigidly from Agent 1 -> Agent 2 -> SOP Engine -> Agent 3 -> Agent 4 -> Agent 5 -> Agent 6.')
    doc.add_paragraph('Retries & Timeouts: Retries are managed inside the LLMClient (with up to 3 fallback provider attempts). Node timeouts are strictly set at 3 seconds per node execution.')
    doc.add_paragraph('\n[PLACEHOLDER: Insert Orchestration graph (LangGraph/DAG) diagram here]\n')
    
    doc.add_heading('5.4 Human-in-the-Loop (HITL) Checkpoints', level=2)
    doc.add_paragraph('CuraBot operates as a supervised agent (L3) with strict HITL intersections.')
    doc.add_paragraph('High Risk Checkpoints: SOP-010 (Red Flag Scanner) and SOP-003 (FAST Stroke) run immediately. If a critical condition is flagged, the orchestrator halts the LLM pipeline completely and triggers an immediate Emergency Override Alert to the user.', style='List Bullet')
    doc.add_paragraph('Low Confidence Checkpoints: If top hypothesis confidence remains <75% after 10 iterations, the system forces a safe conclusion requesting human review.', style='List Bullet')
    doc.add_paragraph('HITL Operating Model: The final output is an "evidence pack" (Diagnostic Report HTML/JSON) built specifically for an approval UI, where clinical mentors review the AI’s reasoning and sign off manually.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 5. AGENTIC DESIGN CONT. (PAGE 11)
    # ---------------------------------------------------------
    doc.add_heading('5.5 Agent-to-Agent Communication Pattern', level=2)
    doc.add_paragraph('Communication Pattern: Shared Memory / Orchestrator Routing.')
    doc.add_paragraph('Agents do not communicate via direct API calls to each other. Instead, the entire session context is persisted in the `DiagnosticState` Pydantic payload (acting as a shared memory event bus). Each agent reads the state, performs its specific analytical task, mutates the state, and the orchestrator routes it to the next node.')
    doc.add_paragraph('Cost & Latency Optimization: To ensure the system meets its <3s latency KPI, only Agents 1, 2, and 5 invoke the LLM. Agents 3, 4, and 6 use high-performance Python rule-based logic to save API costs and reduce processing time.')
    doc.add_paragraph()
    
    doc.add_heading('5.6 Agent Release/Versioning Strategy', level=2)
    doc.add_paragraph('Workflow Versioning: The orchestration LangGraph structure is versioned in Git (e.g., pipeline v2.0). Changes to agent sequences require full regression testing against the 15 Automated SOP Tests.', style='List Bullet')
    doc.add_paragraph('Prompt Pack Versioning: System prompts are stored as constants and versioned in the prompt registry (e.g., prompts v1.2).', style='List Bullet')
    doc.add_paragraph('Policy Versioning: SOP guidelines are versioned according to authorized medical society updates (e.g., AHA 2026 guidelines) and applied via specific policy version tags.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # CONCLUSION & SIGN-OFF (PAGE 12)
    # ---------------------------------------------------------
    doc.add_heading('6. Conclusion & Executive Sign-off', level=1)
    doc.add_paragraph('The CuraBot Agentic AI Differential Diagnosis Tutor is designed with an uncompromising focus on clinical safety, educational efficacy, and enterprise-grade architecture.')
    doc.add_paragraph('By implementing a hybrid model of state-of-the-art LLMs (Gemini, Groq) paired with rigid, rule-based Python logic (15 Clinical SOPs and Bayesian confidence algorithms), CuraBot achieves a highly reliable, zero-hallucination-tolerant environment.')
    doc.add_paragraph('This document verifies that all architectural, security, and functional boundaries are fully defined and ready for clinical sandbox deployment.')
    doc.add_paragraph()
    
    doc.add_heading('Signatures', level=2)
    add_detailed_table(doc, ['Role', 'Name', 'Date', 'Signature'], [
        ['Lead Architect', '_________________', '___________', '_________________'],
        ['Security Lead', '_________________', '___________', '_________________'],
        ['Compliance Officer', '_________________', '___________', '_________________'],
        ['Product Owner', '_________________', '___________', '_________________']
    ])

    # Save the document
    output_path = os.path.join(os.getcwd(), 'CuraBot_12Page_Exact_Template_LLD.docx')
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == '__main__':
    create_12page_exact_lld()
