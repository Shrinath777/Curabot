import os
from docx import Document

def add_table_from_data(doc, data):
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(data[0]):
        hdr_cells[i].text = col
    for row_data in data[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

doc = Document()
doc.add_heading('LLD: Document Control, Metadata & Scope', 0)
doc.add_heading('🏥 CuraBot - Agentic AI Differential Diagnosis Tutor', 1)

doc.add_heading('1. Document Control & Metadata', 2)
add_table_from_data(doc, [
    ['Field', 'Details'],
    ['Project Name', 'CuraBot'],
    ['System Name', 'Agentic AI Differential Diagnosis Tutor'],
    ['Environment', 'Development / Educational'],
    ['Confidentiality', 'Internal / Restricted (Contains Clinical Logic)'],
    ['AI Safety Classification', 'PII/PHI Regulated (Handling simulated patient data)'],
    ['AI Autonomy Level', 'L3: Supervised Agent (Act-with-approval)']
])

doc.add_heading('2. Versioning & Change Log', 2)
add_table_from_data(doc, [
    ['Version', 'Date', 'Author', 'Description', 'Status'],
    ['v0.1', '2026-04-15', 'AI Architect', 'Initial Draft & SOP Framework', 'Draft'],
    ['v0.5', '2026-04-20', 'Security Lead', 'Added Threat Model & Safety Guardrails', 'In Review'],
    ['v1.0', '2026-04-27', 'Project Team', 'Finalized Scope, KPIs, and Approval Matrix', 'Approved']
])

doc.add_heading('3. Approval Matrix', 2)
add_table_from_data(doc, [
    ['Stakeholder Group', 'Required Sign-off', 'Name / Role', 'Date'],
    ['Architecture', 'Yes', 'Lead System Architect', '2026-04-27'],
    ['Security', 'Yes', 'CISO / Security Engineer', '2026-04-27'],
    ['Compliance', 'Yes', 'Regulatory Affairs (Medical AI)', '2026-04-27'],
    ['Product', 'Yes', 'Product Owner', '2026-04-27'],
    ['Ops / DevOps', 'Yes', 'Infrastructure Lead', '2026-04-27']
])

doc.add_heading('4. AI Safety & Data Handling Summary', 2)
doc.add_paragraph('Sensitivity: High. The system processes simulated patient symptoms and clinical hypotheses.', style='List Bullet')
doc.add_paragraph('Data Handling: All data is anonymized. No real-world PII should be ingested into the educational LLM prompts.', style='List Bullet')
doc.add_paragraph('Guardrails: Hard-coded SOP protocols enforce "Emergency Check" and "Human Referral" logic before AI generation.', style='List Bullet')

doc.add_heading('5. Business Context & Scope', 2)
doc.add_heading('🎯 Problem Statement', 3)
doc.add_paragraph('Medical students often struggle with the iterative nature of differential diagnosis. Existing tools are either purely static (textbooks) or lack the "reasoning" feedback loop provided by an attending physician.')
doc.add_paragraph('Pain Points: High manual effort for research, latency in feedback, risk of cognitive biases (anchoring).', style='List Bullet')

doc.add_heading('👥 Stakeholders & Personas', 3)
doc.add_paragraph('Primary Persona: Medical Students - Using the tool to practice diagnostic reasoning.', style='List Bullet')
doc.add_paragraph('Consuming Team: Curriculum Developers - Integrating the bot into clinical simulations.', style='List Bullet')
doc.add_paragraph('Secondary Stakeholders: Clinical Mentors - Reviewing student performance logs.', style='List Bullet')

doc.add_heading('6. Scope Definition', 2)
add_table_from_data(doc, [
    ['In-Scope', 'Out-of-Scope (Non-Goals)'],
    ['Simulated differential diagnosis tutoring', 'Real-time clinical diagnosis for actual patients'],
    ['6-Agent reasoning pipeline (Symptom Normalizer -> Strategy)', 'Prescribing medication or treatments'],
    ['Cognitive bias detection and alerting', 'Integration with live Hospital EMR systems'],
    ['Bayesian confidence level updates', 'Replacement of licensed medical practitioners']
])

doc.add_heading('7. Success Metrics & KPIs', 2)
add_table_from_data(doc, [
    ['KPI', 'Metric Target', 'Baseline', 'Measurement Approach'],
    ['Diagnostic Alignment', '>90% match with Gold Standard cases', '65% (Static tools)', 'Evaluation against 50 curated clinical cases'],
    ['Safety Adherence', '100% Red-Flag detection', '80%', 'Automated unit tests for Emergency SOPs'],
    ['User Engagement', '>15 mins avg. session length', '5 mins', 'WebSocket session logging'],
    ['Latency', '<3s response per agent step', '10s', 'Performance monitoring in dev console']
])

doc.add_heading('8. Assumptions & Constraints', 2)
doc.add_paragraph('Assumptions: Reliable access to LLM APIs (OpenRouter/Gemini); User has basic medical terminology knowledge.', style='List Bullet')
doc.add_paragraph('Constraints: System is limited to the predefined 20 Clinical SOPs; Local compute constraints for React frontend.', style='List Bullet')

doc.add_heading('9. Reference Artifacts', 2)
doc.add_paragraph('BRD: Business Requirements Document', style='List Bullet')
doc.add_paragraph('HLD: High-Level Design / Architecture Diagram', style='List Bullet')
doc.add_paragraph('ADRs: Architectural Decision Records', style='List Bullet')
doc.add_paragraph('Test Plan: Automated Verification Suite', style='List Bullet')

output_path = os.path.join(os.getcwd(), 'LLD_CuraBot_Document.docx')
doc.save(output_path)
print(f"LLD Document successfully created at: {output_path}")
