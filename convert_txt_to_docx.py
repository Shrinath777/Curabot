import os
import re
from docx import Document

def create_doc_from_text(input_file, output_file):
    doc = Document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for i, line in enumerate(lines):
        text = line.strip()
        if not text:
            # Skip empty lines or add blank paragraph? Usually word handles spacing with styles.
            continue
            
        # Detect headings based on numbering
        if re.match(r'^CuraBot — Low-Level Design Document', text):
            doc.add_heading(text, 0)
        elif re.match(r'^\d+\.\s+[A-Za-z]', text):
            # e.g., "1. System Overview" -> Heading 1
            doc.add_heading(text, 1)
        elif re.match(r'^\d+\.\d+\s+[A-Za-z]', text):
            # e.g., "1.1 Problem Statement" -> Heading 2
            doc.add_heading(text, 2)
        elif text.startswith('•'):
            # bullet points
            doc.add_paragraph(text[1:].strip(), style='List Bullet')
        else:
            # Bold certain prefixes if they exist (like "Primary Users:")
            if ":" in text and len(text.split(":")[0]) < 30 and "{" not in text:
                p = doc.add_paragraph()
                parts = text.split(":", 1)
                p.add_run(parts[0] + ":").bold = True
                if parts[1].strip():
                    p.add_run(" " + parts[1].strip())
            else:
                doc.add_paragraph(text)
                
    doc.save(output_file)
    print(f"Document saved to {output_file}")

if __name__ == '__main__':
    input_path = os.path.join(os.getcwd(), 'lld_content.txt')
    output_path = os.path.join(os.getcwd(), 'CuraBot_LLD_Version1.docx')
    create_doc_from_text(input_path, output_path)
