import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_detailed_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

def create_ultimate_lld():
    doc = Document()
    
    # ---------------------------------------------------------
    # TITLE PAGE
    # ---------------------------------------------------------
    title = doc.add_heading('Low Level Design (LLD) Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Enterprise GenAI / Agentic AI Application\nProject: CuraBot — Agentic AI Differential Diagnosis Tutor\nVersion 1.0')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 1. DOCUMENT CONTROL & METADATA
    # ---------------------------------------------------------
    doc.add_heading('1. Document Control & Metadata', level=1)
    
    doc.add_heading('1.1 Document Header', level=2)
    add_detailed_table(doc, ['Project', 'System Name', 'Environment', 'Confidentiality'], [
        ['CuraBot', 'Agentic AI Differential Diagnosis Tutor', 'Dev / Int / UAT / Prod', 'Strictly Internal / Restricted']
    ])
    
    doc.add_heading('1.2 Doc Control Table & Change Log', level=2)
    # The template specifically requested (Version/Author/Reviewer/Approver/Date)
    add_detailed_table(doc, ['Version', 'Author', 'Reviewer', 'Approver', 'Date', 'Change Log'], [
        ['v0.1 draft', 'AI Architect', 'Peer Architect', 'Lead Architect', '2026-04-10', 'Initial Draft'],
        ['v1.0 approved', 'Project Team', 'Security Lead', 'ARB/TRB', '2026-04-27', 'Final Approval']
    ])
    
    doc.add_heading('1.3 Approval Matrix', level=2)
    add_detailed_table(doc, ['Area', 'Reviewer', 'Approver', 'Evidence/Link'], [
        ['Architecture', 'Lead Architect', 'ARB/TRB', 'ADR list & C4 Diagrams'],
        ['Security', 'Security Lead', 'CISO delegate', 'Threat model'],
        ['Compliance', 'Compliance SME', 'AIRB/Compliance board', 'Policy checklist']
    ])
    
    doc.add_heading('1.4 AI Safety Classification & Data Handling Summary', level=2)
    doc.add_paragraph('Safety Classification: PII/PHI Regulated.', style='List Bullet')
    doc.add_paragraph('Data Sensitivity Summary: High. The system handles simulated patient health data. All data is anonymized prior to LLM transmission. Strict tenant isolation ensures no data bleeds between user sessions.', style='List Bullet')
    
    doc.add_heading('1.5 AI Autonomy Level', level=2)
    doc.add_paragraph('Autonomy Level: L3 Supervised Agent.', style='List Bullet')
    doc.add_paragraph('Type: "Recommend-only" and "Act-with-approval". The system does not take final medical actions automatically.', style='List Bullet')
    
    doc.add_heading('1.6 Reference Links', level=2)
    doc.add_paragraph('BRD, HLD, ADRs, Threat Model, and Test Plan are linked in the internal Confluence repository.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 2. SYSTEM OVERVIEW
    # ---------------------------------------------------------
    doc.add_heading('2. System Overview', level=1)
    
    doc.add_heading('2.1 Business Context, Stakeholders & Personas', level=2)
    doc.add_paragraph('Primary Personas: Medical Students (practicing diagnostics) and Patients (checking symptom urgency).', style='List Bullet')
    doc.add_paragraph('Consuming Teams: Curriculum Developers and Clinical Mentors.', style='List Bullet')
    doc.add_paragraph('Primary Use Cases: 1) Multi-turn diagnostic reasoning, 2) Triage Red-Flag scanning, 3) Medical PDF interpretation.', style='List Bullet')
    
    doc.add_heading('2.2 Problem Statement', level=2)
    doc.add_paragraph('Medical students face high manual effort and slow feedback latency when practicing clinical reasoning. The measurable pain includes lack of dynamic feedback and high risk of unchecked cognitive biases.')
    
    doc.add_heading('2.3 Scope Boundaries', level=2)
    add_detailed_table(doc, ['In-Scope User Journeys', 'Out-of-Scope (Explicit Non-Goals)'], [
        ['Symptom Normalization', 'Real-time diagnosis of actual/live patients'],
        ['Generation of Ranked Diagnoses', 'Prescribing medication or treatments'],
        ['RAG for Medical PDFs', 'Integration with live Hospital EMR systems']
    ])
    
    doc.add_heading('2.4 Success Metrics & KPIs', level=2)
    add_detailed_table(doc, ['KPI Category', 'Measurement Approach (Baseline -> Target)'], [
        ['Quality (Diagnostic Alignment)', 'Automated case evaluation (65% -> >90%)'],
        ['Time (Latency)', 'Dev Console telemetry (10s -> <3s per step)'],
        ['Compliance (Safety Adherence)', 'SOP Unit tests (80% -> 100% Red-Flag detection)'],
        ['Adoption (Engagement)', 'Session logging (5 mins -> >15 mins)']
    ])
    
    doc.add_heading('2.5 Assumptions, Constraints & Autonomy Target', level=2)
    doc.add_paragraph('Assumptions: Data availability of disease knowledge base, LLM model access via API.', style='List Bullet')
    doc.add_paragraph('Constraints: Tool access limited to defined SOPs. No external web scraping allowed.', style='List Bullet')
    doc.add_paragraph('Autonomy Level Target: Started as L1 assistant → target achieved is L3 Supervised Agent (Recommended for regulated).', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 3. ARCHITECTURE DECOMPOSITION
    # ---------------------------------------------------------
    doc.add_heading('3. Architecture Decomposition', level=1)
    
    doc.add_heading('3.1 C4-Context', level=2)
    doc.add_paragraph('Actors: End Users, Mentors. Upstream: None. Downstream: Supabase, ChromaDB, LLMs. Trust Boundaries: API Gateway.')
    doc.add_paragraph('[PLACEHOLDER: Insert PlantUML / draw.io C4-Context Diagram Here]')
    
    doc.add_heading('3.2 C4-Container', level=2)
    doc.add_paragraph('Containers: React UI, FastAPI API Gateway, LangGraph Orchestrator, ChromaDB Retrieval, Python Agent Runtime, Supabase Data Stores.')
    doc.add_paragraph('[PLACEHOLDER: Insert PlantUML / draw.io C4-Container Diagram Here]')
    
    doc.add_heading('3.3 Component Inventory with Owners (C4-Component)', level=2)
    add_detailed_table(doc, ['Component', 'Description', 'Owner/Team'], [
        ['Retriever Module', 'ChromaDB interface for RAG', 'Data Eng Team'],
        ['Ranker / Router', 'LLMClient API rotation logic', 'Backend Core Team'],
        ['Prompt Builder', 'Dynamic DiagnosticState injector', 'AI/ML Team'],
        ['Validators (SOPs)', '15 Python guardrail engines', 'Clinical QA Team']
    ])
    
    doc.add_heading('3.4 C4-Code', level=2)
    doc.add_paragraph('Key packages: `services/orchestrator.py`, `services/llm_client.py`. Interfaces: REST JSON.')
    
    doc.add_heading('3.5 Deployment Topology & Environments', level=2)
    doc.add_paragraph('Topology: K8s cluster, namespaces (curabot-ui, curabot-api). Network zones/VNETs for isolation. Private endpoints for Supabase.')
    doc.add_paragraph('Environments: Dev, Int, UAT, Prod. Config strategy: .env secrets and feature flags (LaunchDarkly) for LLM routing.')
    doc.add_paragraph('[PLACEHOLDER: Insert Deployment Diagram + Network Zones Here]')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 4. GENAI CAPABILITY DESIGN
    # ---------------------------------------------------------
    doc.add_heading('4. GenAI Capability Design', level=1)
    
    doc.add_heading('4.1 RAG Pipeline Flow & Retrieval Approach', level=2)
    doc.add_paragraph('Approach: Vector-RAG (combining vector embeddings + user metadata scoping).')
    doc.add_paragraph('RAG Flow Step-by-step:', style='List Bullet')
    doc.add_paragraph('1. Ingest: pdfplumber extracts text.', style='List Bullet 2')
    doc.add_paragraph('2. Index: 600-char chunks embedded via GoogleGenerativeAiEmbeddingFunction.', style='List Bullet 2')
    doc.add_paragraph('3. Retrieve: Cosine similarity search scoped by user_id tenant rules.', style='List Bullet 2')
    doc.add_paragraph('4. Rerank: Top-K (3) chunks selected.', style='List Bullet 2')
    doc.add_paragraph('5. Generate: Injected into Agent prompts.', style='List Bullet 2')
    doc.add_paragraph('6. Validate: Output Schema checked.', style='List Bullet 2')
    
    doc.add_heading('4.2 Scoping & Chunking Rules', level=2)
    doc.add_paragraph('Scoping rules: Strict tenant/user isolation. Users cannot query other LOB or user documents.')
    doc.add_paragraph('Chunking strategy: 600-character chunks to preserve clinical line breaks.')
    
    doc.add_heading('4.3 Prompt Registry & Lifecycle', level=2)
    add_detailed_table(doc, ['Prompt Type', 'Version', 'Approval Status', 'Purpose'], [
        ['System Prompts', 'v1.2', 'Approved', 'Defines agent personas and rules'],
        ['Task Prompts', 'v1.0', 'Approved', 'Dynamic state injection templates'],
        ['Retrieval Templates', 'v1.1', 'Approved', 'Formats chunks for context window'],
        ['Evaluation Prompts', 'v1.0', 'Approved', 'Self-Critique bias checking']
    ])
    
    doc.add_heading('4.4 Guardrails & Grounding Policy', level=2)
    doc.add_paragraph('Prompt-injection mitigation: Sanitized via rule-based SOP-010 before LLM processing.', style='List Bullet')
    doc.add_paragraph('Output schema validation: strict JSON requirement enforced by backend parsers.', style='List Bullet')
    doc.add_paragraph('Grounding Policy: "Answer only from retrieved evidence and authorized KB." Enforced via MedicalCitationEngine which outputs PubMed links.', style='List Bullet')
    
    doc.add_heading('4.5 Confidence Scoring Approach', level=2)
    doc.add_paragraph('Generation confidence is calculated via Bayesian likelihood ratios applied to the retrieval confidence and evidence matches, with hard minimum thresholds for critical diseases.')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 5. AGENTIC DESIGN
    # ---------------------------------------------------------
    doc.add_heading('5. Agentic Design', level=1)
    
    doc.add_heading('5.1 Agent Inventory Table & RACI', level=2)
    add_detailed_table(doc, ['Agent', 'Type', 'Inputs', 'Outputs', 'Tools Allowed', 'Cannot Do'], [
        ['Symptom Normalizer', 'Extractor', 'User Msg', 'Symptoms JSON', 'LLM', 'Final Action'],
        ['Hypothesis Gen.', 'Analyzer', 'Symptoms', 'Hypothesis JSON', 'KB Search', 'Final Action'],
        ['Evidence Eval.', 'Governance', 'Hypotheses', 'Ledger', 'Rule Checker', 'Final Action'],
        ['Hypothesis Reviser', 'Utility', 'Ledger', 'Revised Scores', 'Bayesian Tool', 'Final Action'],
        ['Diagnostic Strategist', 'Orchestrator', 'State', 'Next Question', 'LLM', 'Prescribe Meds'],
        ['Self-Critique', 'Validation', 'State', 'Bias Flags', 'Rule Checker', 'Final Action']
    ])
    
    doc.add_heading('5.2 Orchestration Graph', level=2)
    doc.add_paragraph('Design: LangGraph state machine. States transition from Normalizer -> Critique. Timeouts set at 3 seconds per node. Retries managed by LLMClient.')
    doc.add_paragraph('[PLACEHOLDER: Insert Orchestration graph (LangGraph/DAG) diagram here]')
    
    doc.add_heading('5.3 Human-in-the-Loop (HITL) Operating Model', level=2)
    doc.add_paragraph('Checkpoints: Low confidence (<75%) triggers 10-turn max timeout. High risk (Red Flags) triggers immediate Emergency Override.')
    doc.add_paragraph('HITL Operating Model: Generates an "evidence pack" (Diagnostic Report HTML/JSON) for approval UI where mentors can sign off on the session.')
    
    doc.add_heading('5.4 Agent-to-Agent Communication', level=2)
    doc.add_paragraph('Communication Pattern: Shared Memory. Agents do not call each other directly; they read/write to the `DiagnosticState` payload routed by the orchestrator.')
    
    doc.add_heading('5.5 Agent Release/Versioning Strategy', level=2)
    doc.add_paragraph('Workflow Version: Managed in Git (e.g., pipeline v2.0).', style='List Bullet')
    doc.add_paragraph('Prompt Pack Version: Managed in prompt registry (e.g., prompts v1.2).', style='List Bullet')
    doc.add_paragraph('Policy Version: SOP guidelines versioned to medical society updates (e.g., AHA 2026 guidelines).', style='List Bullet')

    output_path = os.path.join(os.getcwd(), 'CuraBot_Exact_Template_LLD.docx')
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == '__main__':
    create_ultimate_lld()
