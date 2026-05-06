import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_detailed_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

def create_12page_lld():
    doc = Document()
    
    # ---------------------------------------------------------
    # PAGE 1: TITLE PAGE
    # ---------------------------------------------------------
    for _ in range(5):
        doc.add_paragraph()
        
    title = doc.add_heading('Low Level Design (LLD) Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Enterprise GenAI / Agentic AI Application\nProject: CuraBot — Agentic AI Differential Diagnosis Tutor\nVersion 1.0 (Comprehensive Build)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    
    for _ in range(15):
        doc.add_paragraph()
        
    doc.add_paragraph('Disclaimer: FOR MEDICAL EDUCATION ONLY — Not for clinical use', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 2: 1. DOCUMENT CONTROL & METADATA
    # ---------------------------------------------------------
    doc.add_heading('1. Document Control & Metadata', level=1)
    doc.add_paragraph('Purpose: Make the LLD auditable, versioned, and approved with clear AI safety classification.', style='Intense Quote')
    
    doc.add_heading('1.1 Document Header', level=2)
    add_detailed_table(doc, ['Field', 'Details'], [
        ['Project Name', 'CuraBot'],
        ['System Name', 'Agentic AI Differential Diagnosis Tutor'],
        ['Target Environment', 'Development / Educational (Clinical Sandbox)'],
        ['Confidentiality Level', 'Strictly Internal / Restricted (Contains Proprietary Clinical Logic)'],
        ['Data Classification', 'High / Protected (Handles Simulated Patient Data & Uploaded PDFs)'],
        ['Sponsor', 'Medical Education Innovation Board']
    ])
    doc.add_paragraph()
    
    doc.add_heading('1.2 Versioning & Change Log', level=2)
    add_detailed_table(doc, ['Version', 'Date', 'Author', 'Description', 'Status'], [
        ['v0.1', '2026-04-10', 'AI Architect', 'Initial Draft & SOP Framework Definition', 'Draft'],
        ['v0.5', '2026-04-15', 'Security Lead', 'Added Threat Model & Agent Guardrails', 'In Review'],
        ['v0.8', '2026-04-20', 'Compliance SME', 'HIPAA & PII Data Handling Overlays Added', 'In Review'],
        ['v1.0', '2026-04-27', 'Project Team', 'Finalized Comprehensive Enterprise GenAI Format', 'Approved']
    ])
    doc.add_paragraph()
    
    doc.add_heading('1.3 Approval Matrix', level=2)
    add_detailed_table(doc, ['Area', 'Reviewer', 'Approver', 'Evidence/Link'], [
        ['Architecture', 'Lead Systems Architect', 'Architecture Review Board (ARB)', 'ADR List / Orchestration Graph / C4'],
        ['Security', 'InfoSec Lead', 'Chief Information Security Officer', 'Threat Model / 10 Clinical SOPs Check'],
        ['Compliance', 'Compliance SME', 'Compliance & Ethics Board', 'Medical AI & PII Policy Checklist'],
        ['Data Governance', 'Data Steward', 'Chief Data Officer', 'Data Retention & RAG Vector Policy'],
        ['Product', 'Product Owner', 'VP of Product', 'BRD & KPI Alignment Matrix'],
        ['Operations', 'DevOps Lead', 'Infrastructure Head', 'Deployment Topology / Resource Specs']
    ])
    doc.add_paragraph()
    
    doc.add_heading('1.4 AI Safety Classification & Data Handling', level=2)
    doc.add_paragraph('Safety Classification:', style='List Bullet')
    doc.add_paragraph('PII/PHI Regulated. The system processes simulated patient symptoms, clinical hypotheses, and uploaded medical PDF reports.', style='List Bullet 2')
    doc.add_paragraph('Data Handling Summary:', style='List Bullet')
    doc.add_paragraph('All data is anonymized prior to any LLM transmission. No real-world PII is ingested into the educational LLM prompts. RAG context is strictly isolated per user session (tenant scoping) utilizing ChromaDB namespaces.', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('1.5 AI Autonomy Level', level=2)
    doc.add_paragraph('Autonomy Level: L3 — Supervised Agent ("Act-with-approval" / "Recommend-only").')
    doc.add_paragraph('The system acts exclusively as a conversational diagnostician that recommends conditions and requests human feedback/approval. It NEVER prescribes medication, executes external treatments, or updates live EMR records autonomously. All actions are overseen by the human user.')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 3: 2. SYSTEM OVERVIEW
    # ---------------------------------------------------------
    doc.add_heading('2. System Overview', level=1)
    doc.add_paragraph('Purpose: Set clear scope boundaries so GenAI/agents do not "accidentally" expand behavior.', style='Intense Quote')
    
    doc.add_heading('2.1 Business Context & Stakeholders', level=2)
    doc.add_paragraph('Medical students and early-career healthcare learners often lack access to structured tools that can help them practice clinical reasoning skills outside the classroom. Traditional medical education relies heavily on textbook cases and limited clinical rotations, which do not provide the iterative, feedback-driven diagnostic thinking that real clinical practice demands.')
    doc.add_paragraph('Primary Personas:', style='List Bullet')
    doc.add_paragraph('Medical Students: The core audience. They use the tool to practice differential diagnosis by entering symptoms and observing how the AI reasons through possible conditions, evaluates evidence, and updates hypotheses over multiple rounds of questioning.', style='List Bullet 2')
    doc.add_paragraph('General Users (Patients): Individuals who want a preliminary understanding of their symptoms before visiting a doctor. The system helps them understand the urgency of their condition (Triage).', style='List Bullet 2')
    doc.add_paragraph('Consuming Teams:', style='List Bullet')
    doc.add_paragraph('Curriculum Developers: Integrating the bot into clinical simulations and using performance logs for grading and curriculum enhancement.', style='List Bullet 2')
    doc.add_paragraph('Clinical Mentors: Reviewing diagnostic reports generated by students to provide human-in-the-loop oversight.', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('2.2 Problem Statement & Measurable Pain', level=2)
    doc.add_paragraph('The lack of dynamic, accessible clinical reasoning tools leads to significant skill gaps in diagnostic formulation.')
    doc.add_paragraph('Measurable Pain Points:', style='List Bullet')
    doc.add_paragraph('High manual effort required by faculty to create and grade synthetic clinical cases.', style='List Bullet 2')
    doc.add_paragraph('Latency in feedback loops; students often wait weeks to receive feedback on diagnostic assessments.', style='List Bullet 2')
    doc.add_paragraph('High risk of clinical cognitive biases (e.g., anchoring or premature closure) going uncorrected during formative training years.', style='List Bullet 2')
    doc.add_paragraph('Patients face delayed care due to a lack of immediate, accessible symptom triage.', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('2.3 Scope Boundaries (In-Scope / Out-of-Scope)', level=2)
    add_detailed_table(doc, ['In-Scope User Journeys', 'Out-of-Scope (Explicit Non-Goals)'], [
        ['Symptom Normalization from plain text', 'Real-time diagnosis of actual/live patients'],
        ['Generation of 3-5 Ranked Differential Diagnoses', 'Prescribing medication, procedures, or treatments'],
        ['Multi-turn conversational follow-up questions', 'Integration with live Hospital EMR / EHR systems'],
        ['PDF Medical Report Upload & RAG processing', 'Replacing licensed medical practitioners'],
        ['Emergency Triage & Red-Flag Interception', 'Autonomous booking of specialist appointments'],
        ['Detection of cognitive biases in reasoning', 'Long-term patient health management/tracking'],
        ['Bayesian confidence level updates over time', 'Integration with wearable health devices (IoT)'],
        ['Profile management (History, Allergies, Meds)', 'Automated billing or medical coding']
    ])
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 4: SYSTEM OVERVIEW (CONTINUED)
    # ---------------------------------------------------------
    doc.add_heading('2.4 Success Metrics & SLAs (KPIs)', level=2)
    doc.add_paragraph('To ensure the agentic system delivers value without compromising clinical education safety, the following strict Service Level Agreements (SLAs) and KPIs are monitored.')
    add_detailed_table(doc, ['KPI Category', 'Target Metric', 'Baseline', 'Measurement Approach'], [
        ['Diagnostic Alignment', '>90% match w/ Gold Standard', '65% (Static tools)', 'Automated evaluation against 50 curated cases'],
        ['Safety Adherence', '100% Red-Flag detection', '80%', 'Unit tests targeting SOP-013 & SOP-008'],
        ['User Engagement', '>15 mins average session', '5 mins', 'WebSocket / Supabase session activity logging'],
        ['System Latency', '<3s response per agent step', '10s', 'Telemetry and Dev Console monitoring'],
        ['RAG Accuracy', '>95% precision on extraction', '70%', 'Ground-truth testing on synthetic PDF reports'],
        ['Hallucination Rate', '<1% ungrounded assertions', '5%', 'Medical Citation Engine coverage tests']
    ])
    doc.add_paragraph()
    
    doc.add_heading('2.5 Assumptions and Constraints', level=2)
    doc.add_paragraph('Assumptions:', style='List Bullet')
    doc.add_paragraph('Reliable external API access for LLMs (Gemini, Groq, OpenRouter) is maintained.', style='List Bullet 2')
    doc.add_paragraph('The user possesses a basic understanding of their symptoms and can communicate them effectively in English.', style='List Bullet 2')
    doc.add_paragraph('Vector database queries will not exceed 1000 chunks per patient record.', style='List Bullet 2')
    doc.add_paragraph('Constraints:', style='List Bullet')
    doc.add_paragraph('The system is strictly constrained to the 15 implemented Clinical SOPs (Standard Operating Procedures). Any deviation requires a codebase update.', style='List Bullet 2')
    doc.add_paragraph('Local compute limitations affect Vector DB processing since ChromaDB runs locally.', style='List Bullet 2')
    doc.add_paragraph('API Rate limits imposed by free-tier LLM providers (handled via fallback rotation).', style='List Bullet 2')
    doc.add_paragraph()
    
    # ---------------------------------------------------------
    # PAGE 5: 3. ARCHITECTURE DECOMPOSITION
    # ---------------------------------------------------------
    doc.add_heading('3. Architecture Decomposition', level=1)
    doc.add_paragraph('Purpose: Provide a layered breakdown (C4) plus deployment topology for enterprise runtime.', style='Intense Quote')
    
    doc.add_heading('3.1 C4-Context Diagram Details', level=2)
    doc.add_paragraph('Actors: End Users (Medical Students, Patients), Clinical Mentors / Reviewers.')
    doc.add_paragraph('Upstream Systems: None (System is the primary entry point).')
    doc.add_paragraph('Downstream Systems:')
    doc.add_paragraph('Supabase (PostgreSQL BaaS) for persistent state.', style='List Bullet 2')
    doc.add_paragraph('ChromaDB (Vector DB) for embedding storage.', style='List Bullet 2')
    doc.add_paragraph('LLM API Providers (Gemini, Groq, OpenRouter, SambaNova).', style='List Bullet 2')
    doc.add_paragraph('Trust Boundaries: All client communication terminates at the FastAPI API Gateway over HTTPS. The agent runtime and orchestration occur securely behind the backend firewall. No direct database access is permitted from the frontend.')
    doc.add_paragraph()
    
    doc.add_heading('3.2 C4-Container Details', level=2)
    doc.add_paragraph('The system is distributed across several decoupled containers:')
    add_detailed_table(doc, ['Container', 'Tech Stack', 'Responsibility'], [
        ['User Interface (UI)', 'React, TypeScript, Vite, Tailwind', 'Renders chat, hypothesis cards, trajectory charts.'],
        ['API Gateway', 'FastAPI, Uvicorn, Python', 'Handles REST endpoints (/api/chat, /api/auth).'],
        ['Orchestrator', 'LangGraph Core', 'Coordinates the 6-agent pipeline workflow.'],
        ['Agent Runtime', 'Python Agents (Classes)', 'Executes individual tasks (Normalization, Eval).'],
        ['Data Stores', 'Supabase (Cloud) / SQLite (Local)', 'Stores users, sessions, profiles, chat history.'],
        ['Vector Database', 'ChromaDB', 'Stores embeddings of medical PDFs for RAG.']
    ])
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 6: ARCHITECTURE DECOMPOSITION (CONTINUED)
    # ---------------------------------------------------------
    doc.add_heading('3.3 C4-Component Breakdown', level=2)
    doc.add_paragraph('Within the backend container, the architecture is highly modularized:')
    doc.add_paragraph('LLMClient (services/llm_client.py):', style='List Bullet')
    doc.add_paragraph('Manages provider fallback chain (Gemini -> Groq -> OpenRouter -> SambaNova). Implements a global async lock (`_api_lock`) and applies a 2.0s pacing lock to mitigate API rate limits (429 errors). Automatically rotates API keys if exhausted.', style='List Bullet 2')
    doc.add_paragraph('RAG Retriever Engine (ChromaDB Integration):', style='List Bullet')
    doc.add_paragraph('Extracts text from PDF via `pdfplumber`, splits into 600-character chunks, indexes, and retrieves the top-3 most relevant chunks per user query via cosine similarity.', style='List Bullet 2')
    doc.add_paragraph('Medical Evidence Citation Engine:', style='List Bullet')
    doc.add_paragraph('A pure Python module generating ICD-10 codes, Clinical Authority links (e.g., AHA, CDC, AGA), and PubMed verification links. Ensures zero-hallucination grounding for all diagnoses.', style='List Bullet 2')
    doc.add_paragraph('Clinical SOP Engine:', style='List Bullet')
    doc.add_paragraph('15 distinct rule-based engines (e.g., Triage, Chest Pain, Stroke, Lab interpretation) running synchronously pre- and post-LLM invocation to enforce hard clinical safety limits.', style='List Bullet 2')
    doc.add_paragraph('Patient History Analyzer:', style='List Bullet')
    doc.add_paragraph('Evaluates past diagnoses and recurring conditions for returning users, injecting contextual baselines into the current session.', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('3.4 Deployment Topology & Environment', level=2)
    doc.add_paragraph('Frontend Deployment: Hosted on Vercel / Netlify edge networks, providing global CDN distribution and fast load times.')
    doc.add_paragraph('Backend & Vector DB Deployment: Containerized using Docker. Deployed to Kubernetes (K8s) or Cloud VMs with isolated Virtual Private Cloud (VPC) subnets. The FastAPI backend communicates with ChromaDB via local high-speed network interfaces.')
    doc.add_paragraph('Environment Separation: Strict isolation between Dev, Staging (UAT), and Prod environments. Supabase projects are totally distinct per environment. All secrets, API keys, and configurations are injected via secure Environment Variables (.env config strategy), preventing credential leakage in source control.')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 7: 4. GENAI CAPABILITY DESIGN
    # ---------------------------------------------------------
    doc.add_heading('4. GenAI Capability Design', level=1)
    doc.add_paragraph('Purpose: Design RAG pipelines, prompt lifecycle, and hallucination controls.', style='Intense Quote')
    
    doc.add_heading('4.1 Retrieval Approach (Vector-RAG)', level=2)
    doc.add_paragraph('Retrieval-Augmented Generation (RAG) is utilized to bring a patient\'s uploaded medical records (PDF lab reports, test results) directly into the diagnostic conversation.')
    doc.add_paragraph('Approach: Vector-RAG applied to the `patient_records` collection.')
    doc.add_paragraph('Scoping Rules: Queries are strictly scoped by `user_id` at the database level to ensure tenants (patients) can never retrieve medical data belonging to another user. Access control is enforced at the API route layer before hitting the vector database.')
    doc.add_paragraph('Chunking Strategy: Text is extracted using `pdfplumber` and split into exactly 600-character chunks. This specific size was chosen to preserve line boundaries for clinical parameters (e.g., "Hb: 12.5 g/dL") without truncating critical context.')
    doc.add_paragraph('Embedding Model: Utilizes `GoogleGenerativeAiEmbeddingFunction` when the Gemini API is available, with a zero-vector local fallback mechanism to maintain system stability if the external API goes offline.')
    doc.add_paragraph()
    
    doc.add_heading('4.2 Prompt Engineering Lifecycle', level=2)
    doc.add_paragraph('CuraBot uses structured, version-controlled prompt engineering to ensure consistent, parseable outputs.')
    doc.add_paragraph('System Prompts (Agent Identity):', style='List Bullet')
    doc.add_paragraph('Static blocks defining roles, clinical rules, and SOPs. For example, the Diagnostic Strategist is strictly instructed to follow the OPQRST clinical workflow (Onset, Provocation, Quality, Region, Severity, Timing).', style='List Bullet 2')
    doc.add_paragraph('Task Prompts (Dynamic Context):', style='List Bullet')
    doc.add_paragraph('Template-based prompts filled dynamically with the current `DiagnosticState` payload. This includes the patient message, hypotheses, evidence ledger, RAG context, and iteration count.', style='List Bullet 2')
    doc.add_paragraph('Evaluation Prompts:', style='List Bullet')
    doc.add_paragraph('Used by the Patient History Analyzer to compare current symptoms against recurring conditions stored in Supabase.', style='List Bullet 2')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 8: GENAI CAPABILITY DESIGN (CONTINUED)
    # ---------------------------------------------------------
    doc.add_heading('4.3 Guardrails & Output Validation', level=2)
    doc.add_paragraph('Schema Validation: All LLM outputs are mandated to follow exact JSON schema structures. Downstream fallback parsers handle markdown-wrapped or malformed outputs. If an LLM fails to output valid JSON, it triggers the local rule-based fallback logic.')
    doc.add_paragraph('Prompt Injection Mitigation: Raw user input is constrained. Emergency phrases and red flags are intercepted by the pure-Python SOP-013 (Red Flag Scanner) BEFORE they ever reach the LLM, neutralizing injection attacks designed to bypass triage.')
    doc.add_paragraph('Citation Requirement & Grounding Policy: The Medical Evidence Citation Engine forces every concluded diagnosis to be backed by deterministically mapped ICD-10 codes and PubMed links. The LLM is strictly instructed: "Answer ONLY using retrieved evidence and the authorized disease knowledge base."')
    doc.add_paragraph()
    
    doc.add_heading('4.4 Confidence Scoring & Bayesian Mathematics', level=2)
    doc.add_paragraph('Confidence scores are NOT hallucinated by the LLM. They are deterministically updated by Agent 4 (Hypothesis Reviser) using a Bayesian likelihood ratio approach:')
    doc.add_paragraph('1. Strong supporting evidence multiplies hypothesis confidence by 1.6x.')
    doc.add_paragraph('2. Moderate supporting evidence multiplies by 1.3x.')
    doc.add_paragraph('3. Weak supporting evidence multiplies by 1.1x.')
    doc.add_paragraph('4. Contradicting evidence applies severe reduction multipliers (Strong: 0.5x, Moderate: 0.7x, Weak: 0.9x).')
    doc.add_paragraph('Severity Floors: Critical-severity diseases cannot drop below an 8% baseline unless they have 3+ contradicting evidence items. This mathematical floor ensures deadly conditions are never prematurely dismissed by the system.')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 9: 5. AGENTIC DESIGN
    # ---------------------------------------------------------
    doc.add_heading('5. Agentic Design', level=1)
    doc.add_paragraph('Purpose: Specify agent inventory, roles/skills, orchestration graph, and HITL checkpoints.', style='Intense Quote')
    
    doc.add_heading('5.1 Agent Inventory & RACI Matrix', level=2)
    doc.add_paragraph('CuraBot operates a strict 6-agent sequential pipeline.')
    add_detailed_table(doc, ['Agent', 'Type', 'Inputs', 'Outputs', 'Tools Allowed', 'Cannot Do'], [
        ['1. Symptom Normalizer', 'Extractor', 'Raw text + RAG History', 'Normalized JSON', 'LLM, Keyword Matching', 'Generate diagnosis, Alter context'],
        ['2. Hypothesis Gen.', 'Analyzer', 'Symptoms + Disease KB', 'Ranked Hypotheses', 'KB Search, LLM', 'Evaluate evidence, Ask user'],
        ['3. Evidence Evaluator', 'Governance', 'Hypotheses + Evidence', 'Evidence Ledger', 'Rule-based evaluation', 'Update confidence scores'],
        ['4. Hypothesis Reviser', 'Math/Logic', 'Ledger + Hypotheses', 'Revised Hypotheses', 'Bayesian Multipliers', 'Ask questions, Access LLM'],
        ['5. Diagnostic Strategist', 'Decision', 'Revised Hypotheses', 'Next Q / Conclusion', 'LLM, OPQRST Workflow', 'Prescribe medications'],
        ['6. Self-Critique', 'Auditor', 'State + Iteration count', 'Bias Flags (JSON)', 'Rule-based checking', 'Alter the final diagnosis']
    ])
    doc.add_paragraph()
    
    doc.add_heading('5.2 The 15 Clinical SOPs (Standard Operating Procedures)', level=2)
    doc.add_paragraph('The agents are bookended by exactly 15 pure Python rule-based SOPs. These act as the definitive safety guardrails of the system:')
    doc.add_paragraph('SOP-001 Triage Severity Classification: Analyzes vitals and text to assign Red/Orange/Yellow/Green urgency levels.', style='List Bullet')
    doc.add_paragraph('SOP-002 Chest Pain Protocol (ACS Pathway): Detects crushing pain radiating to the jaw or arm to assess heart attack risk.', style='List Bullet')
    doc.add_paragraph('SOP-003 Stroke Detection (FAST Protocol): Identifies facial drooping, arm weakness, or speech difficulty for immediate neurological routing.', style='List Bullet')
    doc.add_paragraph('SOP-004 Respiratory Distress Protocol: Evaluates shortness of breath against oxygen saturation (SpO2) vitals.', style='List Bullet')
    doc.add_paragraph('SOP-005 Vital Signs Interpreter: Checks submitted patient vitals (BP, HR, Temp) against standard clinical ranges.', style='List Bullet')
    doc.add_paragraph('SOP-006 Infection Risk Protocol: Combines fever, respiratory symptoms, and exposure history to mandate isolation.', style='List Bullet')
    doc.add_paragraph('SOP-007 Sepsis Screening (qSOFA): Identifies altered mental status and tachypnea indicating potential severe systemic infection.', style='List Bullet')
    doc.add_paragraph('SOP-008 Lab Value Interpreter: Parses extracted PDF RAG text to flag critical lab abnormalities (e.g., critically low Hemoglobin).', style='List Bullet')
    doc.add_paragraph('SOP-009 Medication Safety Check: Cross-references patient allergy profiles against currently prescribed medications to avoid adverse reactions.', style='List Bullet')
    doc.add_paragraph('SOP-010 Red Flag Symptom Scanner: Intercepts high-risk phrases like "worst headache of my life" and forces immediate emergency halts.', style='List Bullet')
    doc.add_paragraph('SOP-011 Specialist Routing: Automatically routes the finalized diagnosis to the appropriate medical specialty (e.g., Cardiology, Neurology).', style='List Bullet')
    doc.add_paragraph('SOP-012 Confidence Calibration: Overrides LLM confidence scores by assessing if sufficient iterations and evidence exist to conclude safely.', style='List Bullet')
    doc.add_paragraph('SOP-013 Pediatric/Geriatric Triage Adjustments: Escalates urgency based on extreme age vulnerability (infants/elderly).', style='List Bullet')
    doc.add_paragraph('SOP-014 Pregnancy Safety Protocol: Scans female patients of childbearing age presenting with abdominal pain for ectopic pregnancy risks.', style='List Bullet')
    doc.add_paragraph('SOP-015 Follow-up Protocol: Generates exact timelines for outpatient or emergency follow-up care based on the final diagnosis.', style='List Bullet')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 10: AGENTIC DESIGN (CONTINUED)
    # ---------------------------------------------------------
    doc.add_heading('5.3 Orchestration Graph & State Machine', level=2)
    doc.add_paragraph('Graph Execution: Managed as a sequential state machine (Directed Acyclic Graph) via LangGraph principles. The pipeline runs exactly once per user message and typically completes the diagnostic cycle in 3-10 iterations.')
    doc.add_paragraph('Communication Pattern: Agents do not communicate via direct API calls. They communicate via Shared Memory. The entire session context is persisted in the `DiagnosticState` Pydantic model (event bus pattern). Each agent reads the state, performs its task, mutates the state, and passes it to the next node.')
    doc.add_paragraph('Cost & Latency Optimization: To ensure the system meets its <3s latency KPI, only Agents 1, 2, and 5 invoke the LLM. Agents 3, 4, and 6 use high-performance Python rule-based logic. This hybrid approach saves massive API costs and drastically reduces processing time.')
    doc.add_paragraph()
    
    doc.add_heading('5.4 Human-in-the-Loop (HITL) Checkpoints', level=2)
    doc.add_paragraph('CuraBot operates as a supervised agent (L3) with strict HITL intersections:')
    doc.add_paragraph('Emergency Override Checkpoint: SOP-010 (Red Flag) & SOP-003 (FAST Stroke) run immediately after normalization. If a critical condition is flagged, the orchestrator halts the LLM pipeline completely and triggers a hard-coded Emergency Alert instructing the user to seek immediate care (911/108).', style='List Bullet')
    doc.add_paragraph('End of Session Review: The final output is flagged prominently as "For Medical Education Only". Any diagnostic application must be reviewed by a human mentor or licensed physician. The system generates a highly detailed evidence pack (JSON/HTML) explicitly for this HITL review.', style='List Bullet')
    doc.add_paragraph()
    
    # ---------------------------------------------------------
    # PAGE 11: DEPLOYMENT & VERSIONING
    # ---------------------------------------------------------
    doc.add_heading('5.5 Release & Versioning Strategy', level=2)
    doc.add_paragraph('Workflow Versioning: The orchestration graph is versioned in Git. Changes to agent sequences or State definitions require full regression testing against the 15 Automated SOP Tests.')
    doc.add_paragraph('Prompt Pack Versioning: System prompts are stored as constants and versioned strictly alongside the codebase. Changes to clinical guidelines (e.g., AHA updates for cardiovascular care) require a major version bump in the prompt registry.')
    doc.add_paragraph('Database Migrations: Supabase schemas are version-controlled via Alembic/Supabase CLI migrations.')
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PAGE 12: CONCLUSION & SIGN-OFF
    # ---------------------------------------------------------
    doc.add_heading('6. Conclusion & Executive Sign-off', level=1)
    doc.add_paragraph('The CuraBot Agentic AI Differential Diagnosis Tutor is designed with an uncompromising focus on clinical safety, educational efficacy, and enterprise-grade architecture.')
    doc.add_paragraph('By implementing a hybrid model of state-of-the-art LLMs (Gemini, Groq) paired with rigid, rule-based Python logic (15 Clinical SOPs and Bayesian confidence algorithms), CuraBot achieves a highly reliable, zero-hallucination-tolerant environment.')
    doc.add_paragraph('This document verifies that all architectural, security, and functional boundaries are fully defined and ready for clinical sandbox deployment.')
    doc.add_paragraph()
    
    doc.add_heading('Signatures', level=2)
    add_detailed_table(doc, ['Role', 'Name', 'Date', 'Signature'], [
        ['Lead Architect', '_________________', '___________', '_________________'],
        ['Security Lead', '_________________', '___________', '_________________'],
        ['Compliance Officer', '_________________', '___________', '_________________'],
        ['Product Owner', '_________________', '___________', '_________________']
    ])

    # Save the document
    output_path = os.path.join(os.getcwd(), 'CuraBot_12Page_Enterprise_LLD_Final.docx')
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == '__main__':
    create_12page_lld()
