import sys
import subprocess
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt, Inches

def create_doc():
    doc = Document()
    
    # Title
    doc.add_heading('Low-Level Design (LLD): AI Safety & Architecture Documentation', 0)
    
    # 1. Document Control
    doc.add_heading('1. Document Control', level=1)
    
    doc.add_heading('1.1 Document Header', level=2)
    doc.add_paragraph('Project: CuraBot - Agentic AI Differential Diagnosis Tutor', style='List Bullet')
    doc.add_paragraph('System Name: CuraBot Diagnostic Engine', style='List Bullet')
    doc.add_paragraph('Environment: Development / Staging / Production', style='List Bullet')
    doc.add_paragraph('Confidentiality: Confidential - Internal Use Only', style='List Bullet')
    
    doc.add_heading('1.2 Versioning & Change Log', level=2)
    table1 = doc.add_table(rows=4, cols=5)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    headers = ['Version', 'Author', 'Description of Changes', 'Date', 'Status']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        
    data1 = [
        ['v0.1', 'Architecture Team', 'Initial draft for AI safety and LLD structure', '2026-04-22', 'Draft'],
        ['v0.2', 'Security Lead', 'Added Threat Model and Data Sensitivity details', '2026-04-22', 'Review'],
        ['v1.0', 'Project Sponsor', 'Final approval and sign-off', '2026-04-23', 'Approved']
    ]
    for i, row in enumerate(data1):
        row_cells = table1.rows[i+1].cells
        for j, val in enumerate(row):
            row_cells[j].text = val

    doc.add_heading('1.3 Approval Matrix', level=2)
    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, h in enumerate(['Area', 'Reviewer', 'Approver', 'Evidence/Link']):
        hdr_cells[i].text = h
    data2 = [
        ['Architecture', 'Lead Architect', 'Architecture Review Board (ARB)', 'ADR List'],
        ['Security', 'Security Lead', 'CISO Delegate', 'System Threat Model'],
        ['Compliance', 'Compliance SME', 'Compliance Board / AIRB', 'Policy & Regulatory Checklist'],
        ['Product', 'Product Manager', 'VP of Product', 'Business Requirements Document (BRD)'],
        ['Ops', 'Lead DevOps Engineer', 'Head of Operations', 'Deployment & Test Plan']
    ]
    for i, row in enumerate(data2):
        row_cells = table2.rows[i+1].cells
        for j, val in enumerate(row):
            row_cells[j].text = val

    doc.add_heading('1.4 AI Safety Classification & Data Handling', level=2)
    p = doc.add_paragraph()
    p.add_run('AI Safety Classification: ').bold = True
    p.add_run('Non-Regulated (Medical Education Use Only). (Note: Should the system ever transition to real patient data, this will be escalated to Regulated/High-Risk).')
    p = doc.add_paragraph()
    p.add_run('Data Sensitivity Summary: ').bold = True
    p.add_run('The system currently processes simulated patient cases. It handles Simulated PHI/PII. Strict data sanitization pipelines are in place to ensure no real patient data is ingested or retained.')

    doc.add_heading('1.5 AI Autonomy Level', level=2)
    p = doc.add_paragraph()
    p.add_run('Current Autonomy Level: ').bold = True
    p.add_run('L3 (Supervised Agent)')
    p = doc.add_paragraph()
    p.add_run('Action Paradigm: ').bold = True
    p.add_run('Recommend-Only')
    p = doc.add_paragraph()
    p.add_run('Description: ').bold = True
    p.add_run('The AI generates hypotheses, evaluates evidence, and provides differential diagnoses for educational purposes. It operates strictly in a "recommendation" capacity and does not execute clinical actions, prescribe treatments, or act without explicit human oversight and approval.')

    doc.add_heading('1.6 Key Reference Artifacts', level=2)
    doc.add_paragraph('Business Requirements Document (BRD)', style='List Bullet')
    doc.add_paragraph('High-Level Design (HLD)', style='List Bullet')
    doc.add_paragraph('Architecture Decision Records (ADRs)', style='List Bullet')
    doc.add_paragraph('Threat Model & Security Assessment', style='List Bullet')
    doc.add_paragraph('Quality Assurance & Test Plan', style='List Bullet')

    doc.add_page_break()

    # 2. System Overview
    doc.add_heading('2. System Overview', level=1)

    doc.add_heading('2.1 Business Context and Stakeholders', level=2)
    doc.add_paragraph('The CuraBot Diagnostic Engine is designed to teach clinical reasoning concepts to medical professionals and students through simulated agentic workflows.')
    p = doc.add_paragraph()
    p.add_run('Primary Personas:').bold = True
    doc.add_paragraph('Medical Students (Primary Users)', style='List Bullet 2')
    doc.add_paragraph('Clinical Instructors / Educators (Curriculum Designers)', style='List Bullet 2')
    doc.add_paragraph('System Administrators', style='List Bullet 2')
    p = doc.add_paragraph()
    p.add_run('Consuming Teams: ').bold = True
    p.add_run('University Medical Programs, Corporate Medical Training Departments.')

    doc.add_heading('2.2 Problem Statement', level=2)
    doc.add_paragraph('Medical education requires extensive, scalable, and personalized feedback on clinical reasoning to prevent diagnostic errors.')
    p = doc.add_paragraph()
    p.add_run('Pain Points:').bold = True
    doc.add_paragraph('High latency in receiving instructor feedback on simulated cases.', style='List Bullet 2')
    doc.add_paragraph('Significant manual effort required by instructors to design and monitor dynamic patient scenarios.', style='List Bullet 2')
    doc.add_paragraph('Difficulty in consistently identifying and correcting cognitive biases (e.g., anchoring bias, premature closure) during the learning process.', style='List Bullet 2')

    doc.add_heading('2.3 Scope Boundaries', level=2)
    table3 = doc.add_table(rows=3, cols=3)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    for i, h in enumerate(['Category', 'Included (In-Scope)', 'Excluded (Out-of-Scope)']):
        hdr_cells[i].text = h
    data3 = [
        ['User Journeys', '- Conducting simulated multi-stage diagnostic interviews (Max 10 turns).\n- Tracking supporting/contradicting evidence.\n- Real-time cognitive bias detection & feedback.', '- Actual clinical diagnosis or treatment planning for real patients.\n- Autonomous prescription generation.'],
        ['Integrations', '- 15 standardized Clinical SOPs.\n- Gemini LLM integration for reasoning.', '- Integration with live EHR/EMR systems.\n- Direct ingestion of real-world patient charts.']
    ]
    for i, row in enumerate(data3):
        row_cells = table3.rows[i+1].cells
        for j, val in enumerate(row):
            row_cells[j].text = val

    doc.add_heading('2.4 Success Metrics and KPIs', level=2)
    table4 = doc.add_table(rows=5, cols=3)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    for i, h in enumerate(['Dimension', 'Metric / KPI', 'Baseline Approach']):
        hdr_cells[i].text = h
    data4 = [
        ['Quality', 'Reduction in cognitive bias occurrences during simulated cases.', 'Track bias flags per session over a student\'s semester.'],
        ['Time', 'Agent response latency (Symptom to Hypothesis generation).', 'Maintain end-to-end response time under < 2.5 seconds.'],
        ['Cost / Effort', 'Reduction in instructor hours per simulated case creation.', 'Measure time spent building vs. CuraBot scenarios.'],
        ['Compliance', '100% adherence to Medical Education disclaimers.', 'Automated pre-flight checks on all UI rendering components.']
    ]
    for i, row in enumerate(data4):
        row_cells = table4.rows[i+1].cells
        for j, val in enumerate(row):
            row_cells[j].text = val

    doc.add_heading('2.5 Assumptions and Constraints', level=2)
    p = doc.add_paragraph()
    p.add_run('Assumptions:').bold = True
    doc.add_paragraph('Users are authenticated and understand the purely educational nature of the platform.', style='List Bullet 2')
    doc.add_paragraph('The underlying LLM (Gemini) maintains consistent reasoning capabilities without sudden API deprecations.', style='List Bullet 2')
    
    p = doc.add_paragraph()
    p.add_run('Constraints:').bold = True
    doc.add_paragraph('Strict limitation to a maximum of 10 interaction turns per diagnostic session to control API costs and prevent hallucination drift.', style='List Bullet 2')
    doc.add_paragraph('System must gracefully fallback if the LLM provider experiences downtime or throttling.', style='List Bullet 2')

    doc.add_heading('2.6 AI Autonomy Level Target', level=2)
    doc.add_paragraph('Current Target: L1 Assistant transitioning to L3 Supervised Agent (Recommend-Only).', style='List Bullet')
    doc.add_paragraph('Future Considerations: The system is explicitly designed not to exceed L3 autonomy to comply with safety guardrails regarding medical applications.', style='List Bullet')

    # Ensure docs folder exists
    os.makedirs(r'c:\projects\tcs project\curabot\docs', exist_ok=True)
    
    output_path = r'c:\projects\tcs project\curabot\docs\LLD_AI_Safety_and_Architecture.docx'
    doc.save(output_path)
    print(f"Word document saved successfully at {output_path}")

if __name__ == '__main__':
    create_doc()
